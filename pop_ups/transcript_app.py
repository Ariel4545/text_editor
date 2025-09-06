# transcript_app.py
# Standalone Transcript popup:
# - Sources: YouTube (ID/URL, optional language) and local audio (mp3/wav)
# - Output formats: TXT/SRT/VTT/MD/HTML/JSON/CSV/DOCX; inline timestamp toggle for TXT
# - Performance & resilience: background execution with busy window, retries with backoff, cancel for long tasks
# - UX/UI: uses app.make_pop_ups_window, app.make_rich_textbox, app.place_toolt (forwarded to enhanced builders)
# - Editor interoperability: insert modes (replace/insert/append); insert captions in chosen format
# - Security & privacy: redact toggle; confirmation when exporting/copying unredacted

import os
import re
import time
import json
import csv
from io import StringIO, BytesIO
from threading import Thread, Event
import tkinter as tk
from tkinter import (
	Label, Frame, Button, Entry, Text, BOTH, DISABLED, END,
	filedialog, messagebox, Checkbutton, Radiobutton, StringVar, BooleanVar, IntVar
)

# Optional dependencies; features degrade gracefully if missing
try:
	from youtube_transcript_api import YouTubeTranscriptApi
	YT_API_AVAILABLE = True
except Exception:
	YouTubeTranscriptApi = None  # type: ignore
	YT_API_AVAILABLE = False

try:
	from pydub import AudioSegment
	PYDUB_AVAILABLE = True
except Exception:
	AudioSegment = None  # type: ignore
	PYDUB_AVAILABLE = False

try:
	from speech_recognition import Recognizer, AudioFile
	SR_AVAILABLE = True
except Exception:
	Recognizer = AudioFile = None  # type: ignore
	SR_AVAILABLE = False

try:
	from docx import Document
	from docx.shared import Pt
	DOCX_AVAILABLE = True
except Exception:
	Document = None
	Pt = None
	DOCX_AVAILABLE = False


def open_transcript(app):
	'''
	Open the Transcript popup window (YouTube or local audio).
	Uses app-level UI builder functions (forwarded to the enhanced builders).
	'''

	# --------------------- shared state ---------------------
	redact_pii_var = BooleanVar(value=False)
	show_inline_timestamp_var = BooleanVar(value=True)
	youtube_language_var = StringVar(value='')
	insert_mode_var = StringVar(value='insert')                 # 'replace' | 'insert' | 'append'
	caption_insert_format_var = StringVar(value='TXT')          # 'TXT' | 'SRT' | 'VTT'

	# Local audio flow controls
	stt_language_var = StringVar(value='')                      # e.g., 'en-US', 'es-ES' (empty => backend default)
	segment_audio_var = BooleanVar(value=False)                 # enable fixed-size chunking
	segment_seconds_var = IntVar(value=90)                      # chunk length in seconds

	# Transcription quality and control toggles
	remove_bracketed_var = BooleanVar(value=True)               # remove [Music], [Applause], etc.
	collapse_whitespace_var = BooleanVar(value=True)            # collapse multi-spaces/newlines

	# --------------------- window helpers ---------------------
	def make_popup_window(window_title: str, *, modal: bool = False, name: str | None = None):
		return app.make_pop_ups_window(
			function=open_transcript,
			custom_title=window_title,
			parent=getattr(app, 'root', None),
			modal=modal,
			name=name,
			topmost=False,
		)

	def show_busy_window(title_text: str, disable_widgets=None, cancel_event: Event | None = None):
		busy_root = make_popup_window(title_text, modal=True, name=f'{title_text}_busy')
		busy_label = Label(busy_root, text=title_text + ' in progress...', font='arial 10')
		busy_label.pack(padx=10, pady=8)
		buttons_row = Frame(busy_root)
		buttons_row.pack(pady=(0, 8))
		if cancel_event is not None:
			cancel_button = Button(buttons_row, text='Cancel', bd=1, command=cancel_event.set)
			cancel_button.pack()
			app.place_toolt([(cancel_button, 'Cancel current task')])
		for widget in (disable_widgets or []):
			try:
				widget.config(state='disabled')
			except Exception:
				pass
		return busy_root

	def finish_busy_window(busy_root, disable_widgets=None, callback=None, result=None, error=None):
		try:
			if busy_root and busy_root.winfo_exists():
				busy_root.destroy()
		except Exception:
			pass
		for widget in (disable_widgets or []):
			try:
				widget.config(state='normal')
			except Exception:
				pass
		if callable(callback):
			callback(result, error)

	def run_async(title_text: str, task_fn, on_done=None, disable_widgets=None, cancel_event: Event | None = None):
		busy_root = show_busy_window(title_text, disable_widgets=disable_widgets, cancel_event=cancel_event)

		def worker():
			task_error, task_result = None, None
			try:
				task_result = task_fn()
			except Exception as exc:
				task_error = exc
			(getattr(app, 'root', None) or app).after(
				0, lambda: finish_busy_window(busy_root, disable_widgets, on_done, task_result, task_error)
			)

		Thread(target=worker, daemon=True).start()

	# --------------------- privacy & sanitation ---------------------
	def redact_text(content: str) -> str:
		if not content:
			return content
		parts = content.split()
		for token_index, token_value in enumerate(parts):
			try:
				if '@' in token_value:
					local_part, sep_char, domain_part = token_value.partition('@')
					if sep_char and local_part and domain_part and '.' in domain_part and ' ' not in domain_part:
						parts[token_index] = '[redacted@email]'
			except Exception:
				continue
		sanitized_text = ' '.join(parts)

		def mask_long_numbers(raw_text: str) -> str:
			digits = '0123456789'
			separators = set('+-(). ')
			output_chars = []
			char_index = 0
			text_length = len(raw_text)
			while char_index < text_length:
				if raw_text[char_index] in digits or raw_text[char_index] in separators:
					segment_start = char_index
					digit_count = 0
					while char_index < text_length and (raw_text[char_index] in digits or raw_text[char_index] in separators):
						if raw_text[char_index] in digits:
							digit_count += 1
						char_index += 1
					segment_text = raw_text[segment_start:char_index]
					output_chars.append('[redacted:number]' if digit_count >= 7 else segment_text)
				else:
					output_chars.append(raw_text[char_index])
					char_index += 1
			return ''.join(output_chars)

		sanitized_text = mask_long_numbers(sanitized_text)

		try:
			marker_prefix = 'ghp_'
			search_start = 0
			output_parts = []
			while True:
				found_index = sanitized_text.find(marker_prefix, search_start)
				if found_index == -1:
					output_parts.append(sanitized_text[search_start:])
					break
				output_parts.append(sanitized_text[search_start:found_index])
				end_index = found_index + len(marker_prefix)
				while end_index < len(sanitized_text) and not sanitized_text[end_index].isspace() and (end_index - found_index) < 48:
					end_index += 1
				output_parts.append('[redacted:token]')
				search_start = end_index
			sanitized_text = ''.join(output_parts)
		except Exception:
			pass
		return sanitized_text

	def normalize_transcript_text(raw_text: str) -> str:
		if not raw_text:
			return raw_text
		normalized_text = raw_text
		if remove_bracketed_var.get():
			try:
				normalized_text = re.sub(r'\[[^\]]{0,40}\]', '', normalized_text)
			except Exception:
				pass
		if collapse_whitespace_var.get():
			try:
				normalized_text = re.sub(r'[ \t]+', ' ', normalized_text)
				normalized_text = re.sub(r'\n{3,}', '\n\n', normalized_text)
				normalized_text = normalized_text.strip()
			except Exception:
				pass
		return normalized_text

	def confirm_non_redacted(action_label: str) -> bool:
		if not redact_pii_var.get():
			return messagebox.askyesno(
				getattr(app, 'title_struct', '') + 'transcript',
				f'{action_label} without redaction?\nThis may include personal data.'
			)
		return True

	# --------------------- formatting helpers ---------------------
	def format_time_hhmmss_mmm(seconds_value: float, separator: str) -> str:
		try:
			total_ms = int(round(float(seconds_value) * 1000.0))
			ms = total_ms % 1000
			total_seconds = total_ms // 1000
			hours = total_seconds // 3600
			minutes = (total_seconds % 3600) // 60
			seconds = total_seconds % 60
			return f'{hours:02d}:{minutes:02d}:{seconds:02d}{separator}{ms:03d}'
		except Exception:
			return f'00:00:00{separator}000'

	def normalize_caption_items(caption_items: list) -> list:
		if not caption_items:
			return []
		normalized_list = []
		for item_index, item_data in enumerate(caption_items):
			start_value = float(item_data.get('start', 0.0))
			duration_value = item_data.get('duration', None)
			if duration_value is None:
				try:
					next_start = float(caption_items[item_index + 1].get('start', start_value + 2.0))
					duration_value = max(0.5, next_start - start_value)
				except Exception:
					duration_value = 2.0
			normalized_list.append({
				'start': start_value,
				'duration': float(duration_value),
				'text': normalize_transcript_text(item_data.get('text', '')),
			})
		return normalized_list

	def wrap_caption_line(line_text: str, max_chars_per_line: int = 40) -> str:
		try:
			words = line_text.split(' ')
			lines = []
			current_line = ''
			for word in words:
				candidate = word if current_line == '' else current_line + ' ' + word
				if len(candidate) <= max_chars_per_line:
					current_line = candidate
				else:
					if current_line:
						lines.append(current_line)
					current_line = word
			if current_line:
				lines.append(current_line)
			if len(lines) > 2:
				merged_text = ' '.join(lines)
				lines = []
				current_line = ''
				for word in merged_text.split(' '):
					candidate = word if current_line == '' else current_line + ' ' + word
					if len(candidate) <= max_chars_per_line:
						current_line = candidate
					else:
						lines.append(current_line)
						current_line = word
				if current_line:
					lines.append(current_line)
				lines = lines[:2]
			return '\n'.join(lines)
		except Exception:
			return line_text

	# --------------------- exports (TXT/SRT/VTT/MD/HTML + JSON/CSV/DOCX) ---------------------
	def export_as_txt(data, include_inline_timestamps: bool) -> str:
		if isinstance(data, list):
			lines = []
			for caption in normalize_caption_items(data):
				text_line = caption['text']
				if include_inline_timestamps:
					timecode_text = format_time_hhmmss_mmm(caption['start'], '.')
					lines.append(f'[{timecode_text}] {text_line}')
				else:
					lines.append(text_line)
			content_text = '\n'.join(lines)
			return redact_text(content_text) if redact_pii_var.get() else content_text
		content_text = normalize_transcript_text(str(data or ''))
		return redact_text(content_text) if redact_pii_var.get() else content_text

	def export_as_srt(caption_items: list) -> str:
		blocks = []
		for block_index, caption in enumerate(normalize_caption_items(caption_items), start=1):
			start_tc = format_time_hhmmss_mmm(caption['start'], ',')
			end_tc = format_time_hhmmss_mmm(caption['start'] + caption['duration'], ',')
			body_text = wrap_caption_line(caption.get('text', ''))
			blocks.append(f'{block_index}\n{start_tc} --> {end_tc}\n{body_text}\n')
		content_text = '\n'.join(blocks).strip()
		return redact_text(content_text) if redact_pii_var.get() else content_text

	def export_as_vtt(caption_items: list) -> str:
		lines = ['WEBVTT', '']
		for caption in normalize_caption_items(caption_items):
			start_tc = format_time_hhmmss_mmm(caption['start'], '.')
			end_tc = format_time_hhmmss_mmm(caption['start'] + caption['duration'], '.')
			body_text = wrap_caption_line(caption.get('text', ''))
			lines.append(f'{start_tc} --> {end_tc}')
			lines.append(body_text)
			lines.append('')
		content_text = '\n'.join(lines).strip()
		return redact_text(content_text) if redact_pii_var.get() else content_text

	def export_as_markdown(caption_items_or_text) -> str:
		if isinstance(caption_items_or_text, list):
			lines = ['# Transcript', '']
			for caption in normalize_caption_items(caption_items_or_text):
				timecode_text = format_time_hhmmss_mmm(caption['start'], '.')
				lines.append(f'- [{timecode_text}] {caption["text"]}')
			content_text = '\n'.join(lines)
		else:
			content_text = f'# Transcript\n\n{normalize_transcript_text(str(caption_items_or_text or ""))}'
		return redact_text(content_text) if redact_pii_var.get() else content_text

	def export_as_html(caption_items_or_text) -> str:
		if isinstance(caption_items_or_text, list):
			html_parts = ['<!doctype html>', '<meta charset="utf-8">', '<h1>Transcript</h1>', '<ul>']
			for caption in normalize_caption_items(caption_items_or_text):
				timecode_text = format_time_hhmmss_mmm(caption['start'], '.')
				html_parts.append(f'<li><strong>[{timecode_text}]</strong> {caption["text"]}</li>')
			html_parts.append('</ul>')
			content_text = '\n'.join(html_parts)
		else:
			safe_text = normalize_transcript_text(str(caption_items_or_text or ''))
			content_text = f'<!doctype html>\n<meta charset="utf-8">\n<h1>Transcript</h1>\n<p>{safe_text}</p>'
		return redact_text(content_text) if redact_pii_var.get() else content_text

	def export_as_timestamp_links(video_id: str, caption_items: list) -> str:
		normalized_id = normalize_youtube_input(video_id)
		base_url = f'https://youtu.be/{normalized_id}'
		links_list = []
		for caption in normalize_caption_items(caption_items):
			start_seconds = int(caption['start'])
			minutes_value = start_seconds // 60
			seconds_value = start_seconds % 60
			absolute_seconds = minutes_value * 60 + seconds_value
			links_list.append(f'{base_url}?t={absolute_seconds}  [{minutes_value:02d}:{seconds_value:02d}] {caption["text"]}')
		content_text = '\n'.join(links_list)
		return redact_text(content_text) if redact_pii_var.get() else content_text

	def export_as_json(caption_items: list) -> str:
		rows = []
		for caption in normalize_caption_items(caption_items):
			rows.append({
				'start': float(caption['start']),
				'end': float(caption['start'] + caption['duration']),
				'duration': float(caption['duration']),
				'text': caption.get('text', ''),
			})
		content_text = json.dumps(rows, ensure_ascii=False, indent=2)
		return redact_text(content_text) if redact_pii_var.get() else content_text

	def export_as_csv(caption_items: list) -> str:
		buffer_obj = StringIO()
		csv_writer = csv.writer(buffer_obj)
		csv_writer.writerow(['start', 'end', 'duration', 'text'])
		for caption in normalize_caption_items(caption_items):
			start_val = float(caption['start'])
			end_val = float(caption['start'] + caption['duration'])
			row_text = caption.get('text', '')
			if redact_pii_var.get():
				row_text = redact_text(row_text)
			csv_writer.writerow([f'{start_val:.3f}', f'{end_val:.3f}', f'{float(caption["duration"]):.3f}', row_text])
		return buffer_obj.getvalue()

	def export_as_docx_from_captions(title_text: str, caption_items: list) -> bytes:
		if not DOCX_AVAILABLE or Document is None or Pt is None:
			messagebox.showerror(getattr(app, 'title_struct', '') + 'transcript', 'python-docx is not available.')
			return b''
		document_obj = Document()
		heading_paragraph = document_obj.add_heading(title_text, level=1)
		for run_item in heading_paragraph.runs:
			try:
				run_item.font.size = Pt(18)
			except Exception:
				pass
		for caption in normalize_caption_items(caption_items):
			paragraph_obj = document_obj.add_paragraph()
			timecode_text = format_time_hhmmss_mmm(caption['start'], '.')
			timecode_run = paragraph_obj.add_run(f'[{timecode_text}] ')
			try:
				timecode_run.italic = True
			except Exception:
				pass
			paragraph_obj.add_run(caption.get('text', ''))
		bytes_buffer = BytesIO()
		document_obj.save(bytes_buffer)
		return bytes_buffer.getvalue()

	def export_as_docx_from_text(title_text: str, text_value: str) -> bytes:
		if not DOCX_AVAILABLE or Document is None or Pt is None:
			messagebox.showerror(getattr(app, 'title_struct', '') + 'transcript', 'python-docx is not available.')
			return b''
		document_obj = Document()
		heading_paragraph = document_obj.add_heading(title_text, level=1)
		for run_item in heading_paragraph.runs:
			try:
				run_item.font.size = Pt(18)
			except Exception:
				pass
		for line_value in (text_value or '').splitlines():
			document_obj.add_paragraph(line_value)
		bytes_buffer = BytesIO()
		document_obj.save(bytes_buffer)
		return bytes_buffer.getvalue()

	def save_text_with_dialog(default_extension: str, content_text: str, default_name: str = 'transcript'):
		save_path = filedialog.asksaveasfilename(
			title='Save transcript',
			defaultextension=default_extension,
			filetypes=[
				('Text', '*.txt'),
				('SubRip', '*.srt'),
				('WebVTT', '*.vtt'),
				('Markdown', '*.md'),
				('HTML', '*.html'),
				('JSON', '*.json'),
				('CSV', '*.csv'),
				('All files', '*.*'),
			],
			initialfile=f'{default_name}{default_extension}',
		)
		if not save_path:
			return
		with open(save_path, 'w', encoding='utf-8') as output_fp:
			output_fp.write(content_text)
		try:
			messagebox.showinfo(getattr(app, 'title_struct', '') + 'transcript', 'Saved successfully.')
		except Exception:
			pass

	def save_bytes_with_dialog(default_extension: str, content_bytes: bytes, default_name: str = 'transcript'):
		save_path = filedialog.asksaveasfilename(
			title='Save transcript',
			defaultextension=default_extension,
			filetypes=[('DOCX', '*.docx'), ('All files', '*.*')],
			initialfile=f'{default_name}{default_extension}',
		)
		if not save_path:
			return
		if not isinstance(content_bytes, (bytes, bytearray)) or len(content_bytes) == 0:
			return
		with open(save_path, 'wb') as output_fp:
			output_fp.write(content_bytes)
		try:
			messagebox.showinfo(getattr(app, 'title_struct', '') + 'transcript', 'Saved successfully.')
		except Exception:
			pass

	# --------------------- editor interoperability ---------------------
	def insert_into_editor(content_text: str, insert_mode: str):
		editor_widget = getattr(app, 'EgonTE', None)
		if editor_widget is None:
			messagebox.showinfo(getattr(app, 'title_struct', '') + 'transcript', 'Editor is not available.')
			return
		text_to_insert = redact_text(content_text) if redact_pii_var.get() else content_text
		if insert_mode == 'replace':
			try:
				editor_widget.delete('sel.first', 'sel.last')
				editor_widget.insert('insert', text_to_insert)
			except Exception:
				editor_widget.insert('insert', text_to_insert)
		elif insert_mode == 'append':
			editor_widget.insert('end', ('\n' if str(editor_widget.get('end-2c')) != '\n' else '') + text_to_insert)
		else:
			editor_widget.insert('insert', text_to_insert)
		try:
			editor_widget.see('insert')
			editor_widget.focus_set()
		except Exception:
			pass

	# --------------------- YouTube helpers ---------------------
	def normalize_youtube_input(raw_input: str) -> str:
		text_value = (raw_input or '').strip()
		if 'youtube.com' in text_value or 'youtu.be' in text_value:
			if 'watch?v=' in text_value:
				text_value = text_value.split('watch?v=', 1)[1]
			elif 'youtu.be/' in text_value:
				text_value = text_value.split('youtu.be/', 1)[1]
			text_value = text_value.split('&', 1)[0]
		return text_value

	def looks_like_youtube_id(candidate_text: str) -> bool:
		if not candidate_text:
			return False
		youtube_id = candidate_text.strip()
		if len(youtube_id) < 6 or len(youtube_id) > 64:
			return False
		for char_value in youtube_id:
			if not (char_value.isalnum() or char_value in ['-', '_']):
				return False
		return True

	def fetch_youtube_captions(video_id_or_url: str, language_code: str | None) -> list:
		if not YT_API_AVAILABLE or YouTubeTranscriptApi is None:
			raise RuntimeError('YouTube transcript API is not available.')
		normalized_id = normalize_youtube_input(video_id_or_url)
		if not looks_like_youtube_id(normalized_id):
			raise ValueError('Invalid YouTube ID/URL.')
		captions_data = None
		chosen_language = (language_code or '').strip().lower()
		if chosen_language:
			transcript_list = YouTubeTranscriptApi.list_transcripts(normalized_id)
			try:
				transcript_obj = transcript_list.find_transcript([chosen_language])
				captions_data = transcript_obj.fetch()
			except Exception:
				for transcript_obj in transcript_list:
					try:
						if getattr(transcript_obj, 'is_translatable', False):
							captions_data = transcript_obj.translate(chosen_language).fetch()
							break
					except Exception:
						continue
		if captions_data is None:
			captions_data = YouTubeTranscriptApi.get_transcript(normalized_id)
		return [{
			'start': float(item.get('start', 0.0)),
			'duration': float(item.get('duration', 2.0)),
			'text': item.get('text', ''),
		} for item in captions_data]

	# --------------------- result windows ---------------------
	def render_text_result_window(window_title: str, plain_text: str):
		result_root = make_popup_window(window_title, name='transcript_text_result')

		final_text = export_as_txt(plain_text, include_inline_timestamps=False)

		container_frame, text_widget, scrollbar_widget = app.make_rich_textbox(root=result_root, place='pack_top')
		container_frame
		scrollbar_widget
		text_widget.insert('1.0', final_text)
		text_widget.configure(state=DISABLED)

		options_frame = Frame(result_root)
		options_frame.pack(pady=6)
		insert_mode_label = Label(options_frame, text='Insert mode:')
		insert_mode_label.grid(row=0, column=0, padx=(0, 6))
		insert_replace_radio = Radiobutton(options_frame, text='Replace selection', variable=insert_mode_var, value='replace')
		insert_cursor_radio = Radiobutton(options_frame, text='Insert at cursor', variable=insert_mode_var, value='insert')
		insert_append_radio = Radiobutton(options_frame, text='Append to end', variable=insert_mode_var, value='append')
		insert_replace_radio.grid(row=0, column=1, padx=6, sticky='w')
		insert_cursor_radio.grid(row=0, column=2, padx=6, sticky='w')
		insert_append_radio.grid(row=0, column=3, padx=6, sticky='w')

		button_row = Frame(result_root)
		button_row.pack(pady=6)

		def on_copy_to_clipboard():
			if not confirm_non_redacted('Copy'):
				return
			to_copy_text = export_as_txt(plain_text, include_inline_timestamps=False)
			result_root.clipboard_clear()
			result_root.clipboard_append(to_copy_text)

		def on_save_txt():
			if not confirm_non_redacted('Save'):
				return
			save_text_with_dialog('.txt', export_as_txt(plain_text, include_inline_timestamps=False))

		def on_save_md():
			if not confirm_non_redacted('Save'):
				return
			save_text_with_dialog('.md', export_as_markdown(plain_text))

		def on_save_html():
			if not confirm_non_redacted('Save'):
				return
			save_text_with_dialog('.html', export_as_html(plain_text))

		def on_save_docx_text():
			if not confirm_non_redacted('Save'):
				return
			doc_text = export_as_txt(plain_text, include_inline_timestamps=False)
			doc_bytes = export_as_docx_from_text('Transcript', doc_text)
			if not isinstance(doc_bytes, (bytes, bytearray)) or len(doc_bytes) == 0:
				return
			save_bytes_with_dialog('.docx', doc_bytes)

		copy_button = Button(button_row, text='Copy', bd=1, command=on_copy_to_clipboard)
		insert_button = Button(button_row, text='Insert to editor', bd=1,
							   command=lambda: insert_into_editor(export_as_txt(plain_text, False),
																  insert_mode_var.get()))
		save_txt_button = Button(button_row, text='Save TXT', bd=1, command=on_save_txt)
		save_md_button = Button(button_row, text='Save MD', bd=1, command=on_save_md)
		save_html_button = Button(button_row, text='Save HTML', bd=1, command=on_save_html)
		save_docx_button = Button(button_row, text='Save DOCX', bd=1, command=on_save_docx_text)
		copy_button.grid(row=0, column=0, padx=6)
		insert_button.grid(row=0, column=1, padx=6)
		save_txt_button.grid(row=0, column=2, padx=6)
		save_md_button.grid(row=0, column=3, padx=6)
		save_html_button.grid(row=0, column=4, padx=6)
		save_docx_button.grid(row=0, column=5, padx=6)

		app.place_toolt([
			(copy_button, 'Copy transcript to clipboard'),
			(insert_button, 'Insert transcript into the editor'),
			(save_txt_button, 'Save as plain text'),
			(save_md_button, 'Save as Markdown'),
			(save_html_button, 'Save as HTML'),
			(save_docx_button, 'Save as DOCX'),
		])

		privacy_note_label = Label(result_root, text='Note: Transcripts may include personal data.', font='arial 8', fg='grey40')
		privacy_note_label.pack(pady=(0, 6))

		try:
			result_root.bind('<Control-c>', lambda event: (on_copy_to_clipboard(), 'break'))
			result_root.bind('<Control-s>', lambda event: (on_save_txt(), 'break'))
			result_root.bind('<Escape>', lambda event: (result_root.destroy(), 'break'))
		except Exception:
			pass

	def render_caption_result_window(window_title: str, caption_items: list, source_video_id: str = ''):
		result_root = make_popup_window(window_title, name='transcript_caption_result')

		txt_preview = export_as_txt(caption_items, include_inline_timestamps=show_inline_timestamp_var.get())

		container_frame, text_widget, scrollbar_widget = app.make_rich_textbox(root=result_root, place='pack_top')
		container_frame
		scrollbar_widget
		text_widget.insert('1.0', txt_preview)
		text_widget.configure(state=DISABLED)

		options_frame = Frame(result_root)
		options_frame.pack(pady=(6, 0))
		show_ts_check = Checkbutton(
			options_frame,
			text='Show inline timestamps (TXT)',
			variable=show_inline_timestamp_var,
			command=lambda: refresh_txt_preview()
		)
		show_ts_check.grid(row=0, column=0, sticky='w', padx=4)

		insert_as_label = Label(options_frame, text='Insert as:')
		insert_as_label.grid(row=0, column=1, padx=(12, 6))
		insert_txt_radio = Radiobutton(options_frame, text='TXT', variable=caption_insert_format_var, value='TXT')
		insert_srt_radio = Radiobutton(options_frame, text='SRT', variable=caption_insert_format_var, value='SRT')
		insert_vtt_radio = Radiobutton(options_frame, text='VTT', variable=caption_insert_format_var, value='VTT')
		insert_txt_radio.grid(row=0, column=2, sticky='w')
		insert_srt_radio.grid(row=0, column=3, sticky='w')
		insert_vtt_radio.grid(row=0, column=4, sticky='w')

		insert_mode_label = Label(options_frame, text='Insert mode:')
		insert_mode_label.grid(row=1, column=0, padx=(0, 6), pady=(6, 0), sticky='w')
		insert_replace_radio = Radiobutton(options_frame, text='Replace selection', variable=insert_mode_var, value='replace')
		insert_cursor_radio = Radiobutton(options_frame, text='Insert at cursor', variable=insert_mode_var, value='insert')
		insert_append_radio = Radiobutton(options_frame, text='Append to end', variable=insert_mode_var, value='append')
		insert_replace_radio.grid(row=1, column=1, padx=6, sticky='w', pady=(6, 0))
		insert_cursor_radio.grid(row=1, column=2, padx=6, sticky='w', pady=(6, 0))
		insert_append_radio.grid(row=1, column=3, padx=6, sticky='w', pady=(6, 0))

		def refresh_txt_preview():
			try:
				text_widget.configure(state='normal')
				text_widget.delete('1.0', 'end')
				new_preview = export_as_txt(caption_items, include_inline_timestamps=show_inline_timestamp_var.get())
				text_widget.insert('1.0', new_preview)
				text_widget.configure(state=DISABLED)
			except Exception:
				pass

		buttons_row = Frame(result_root)
		buttons_row.pack(pady=8)

		def on_copy_captions_txt():
			if not confirm_non_redacted('Copy'):
				return
			content_text = export_as_txt(caption_items, include_inline_timestamps=show_inline_timestamp_var.get())
			result_root.clipboard_clear()
			result_root.clipboard_append(content_text)

		def on_copy_links():
			if not confirm_non_redacted('Copy'):
				return
			links_text = export_as_timestamp_links(source_video_id, caption_items) if source_video_id else export_as_txt(caption_items, True)
			result_root.clipboard_clear()
			result_root.clipboard_append(links_text)

		def on_insert_captions():
			chosen_format = caption_insert_format_var.get()
			if chosen_format == 'SRT':
				content_text = export_as_srt(caption_items)
			elif chosen_format == 'VTT':
				content_text = export_as_vtt(caption_items)
			else:
				content_text = export_as_txt(caption_items, include_inline_timestamps=show_inline_timestamp_var.get())
			insert_into_editor(content_text, insert_mode_var.get())

		def on_save_txt_captions():
			if not confirm_non_redacted('Save'):
				return
			content_text = export_as_txt(caption_items, include_inline_timestamps=show_inline_timestamp_var.get())
			save_text_with_dialog('.txt', content_text)

		def on_save_srt_captions():
			if not confirm_non_redacted('Save'):
				return
			content_text = export_as_srt(caption_items)
			save_text_with_dialog('.srt', content_text)

		def on_save_vtt_captions():
			if not confirm_non_redacted('Save'):
				return
			content_text = export_as_vtt(caption_items)
			save_text_with_dialog('.vtt', content_text)

		def on_save_md_captions():
			if not confirm_non_redacted('Save'):
				return
			content_text = export_as_markdown(caption_items)
			save_text_with_dialog('.md', content_text)

		def on_save_html_captions():
			if not confirm_non_redacted('Save'):
				return
			content_text = export_as_html(caption_items)
			save_text_with_dialog('.html', content_text)

		def on_save_json_captions():
			if not confirm_non_redacted('Save'):
				return
			content_text = export_as_json(caption_items)
			save_text_with_dialog('.json', content_text)

		def on_save_csv_captions():
			if not confirm_non_redacted('Save'):
				return
			content_text = export_as_csv(caption_items)
			save_text_with_dialog('.csv', content_text)

		def on_save_docx_captions():
			if not confirm_non_redacted('Save'):
				return
			doc_bytes = export_as_docx_from_captions('Transcript', caption_items)
			if not isinstance(doc_bytes, (bytes, bytearray)) or len(doc_bytes) == 0:
				return
			save_bytes_with_dialog('.docx', doc_bytes)

		copy_txt_button = Button(buttons_row, text='Copy (TXT)', bd=1, command=on_copy_captions_txt)
		copy_links_button = Button(buttons_row, text='Copy links', bd=1, command=on_copy_links)
		insert_editor_button = Button(buttons_row, text='Insert to editor', bd=1, command=on_insert_captions)
		save_txt_button = Button(buttons_row, text='Save TXT', bd=1, command=on_save_txt_captions)
		save_srt_button = Button(buttons_row, text='Save SRT', bd=1, command=on_save_srt_captions)
		save_vtt_button = Button(buttons_row, text='Save VTT', bd=1, command=on_save_vtt_captions)
		save_md_button = Button(buttons_row, text='Save MD', bd=1, command=on_save_md_captions)
		save_html_button = Button(buttons_row, text='Save HTML', bd=1, command=on_save_html_captions)
		save_json_button = Button(buttons_row, text='Save JSON', bd=1, command=on_save_json_captions)
		save_csv_button = Button(buttons_row, text='Save CSV', bd=1, command=on_save_csv_captions)
		save_docx_button = Button(buttons_row, text='Save DOCX', bd=1, command=on_save_docx_captions)

		copy_txt_button.grid(row=0, column=0, padx=6)
		copy_links_button.grid(row=0, column=1, padx=6)
		insert_editor_button.grid(row=0, column=2, padx=6)
		save_txt_button.grid(row=0, column=3, padx=6)
		save_srt_button.grid(row=0, column=4, padx=6)
		save_vtt_button.grid(row=0, column=5, padx=6)
		save_md_button.grid(row=0, column=6, padx=6)
		save_html_button.grid(row=0, column=7, padx=6)
		save_json_button = Button(buttons_row, text='Save JSON', bd=1, command=on_save_json_captions)
		save_csv_button = Button(buttons_row, text='Save CSV', bd=1, command=on_save_csv_captions)
		save_docx_button = Button(buttons_row, text='Save DOCX', bd=1, command=on_save_docx_captions)

		copy_txt_button.grid(row=0, column=0, padx=6)
		copy_links_button.grid(row=0, column=1, padx=6)
		insert_editor_button.grid(row=0, column=2, padx=6)
		save_txt_button.grid(row=0, column=3, padx=6)
		save_srt_button.grid(row=0, column=4, padx=6)
		save_vtt_button.grid(row=0, column=5, padx=6)
		save_md_button.grid(row=0, column=6, padx=6)
		save_html_button.grid(row=0, column=7, padx=6)
		# realign the second row to start at column 0
		save_json_button.grid(row=1, column=0, padx=6, pady=(6, 0))
		save_csv_button.grid(row=1, column=1, padx=6, pady=(6, 0))
		save_docx_button.grid(row=1, column=2, padx=6, pady=(6, 0))

		app.place_toolt([
			(copy_txt_button, 'Copy captions (TXT)'),
			(copy_links_button, 'Copy timestamp links'),
			(insert_editor_button, 'Insert into editor'),
			(save_txt_button, 'Save as TXT'),
			(save_srt_button, 'Save as SRT'),
			(save_vtt_button, 'Save as VTT'),
			(save_md_button, 'Save as Markdown'),
			(save_html_button, 'Save as HTML'),
			(save_json_button, 'Save as JSON'),
			(save_csv_button, 'Save as CSV'),
			(save_docx_button, 'Save as DOCX'),
			(show_ts_check, 'Toggle inline timestamps preview for TXT'),
			# added helpful hints for radios
			(insert_replace_radio, 'Replace the current selection'),
			(insert_cursor_radio, 'Insert at the caret position'),
			(insert_append_radio, 'Append to the end of the document'),
			(insert_txt_radio, 'Insert captions as plain text'),
			(insert_srt_radio, 'Insert captions formatted as SRT'),
			(insert_vtt_radio, 'Insert captions formatted as WebVTT'),
		])


		privacy_note_label = Label(result_root, text='Note: Transcripts may include personal data.', font='arial 8', fg='grey40')
		privacy_note_label.pack(pady=(0, 6))

		try:
			result_root.bind('<Control-c>', lambda event: (on_copy_captions_txt(), 'break'))
			result_root.bind('<Control-s>', lambda event: (on_save_txt_captions(), 'break'))
			result_root.bind('<Escape>', lambda event: (result_root.destroy(), 'break'))
		except Exception:
			pass

	# --------------------- local audio transcription ---------------------
	def transcribe_local_audio_file():
		chosen_file_path = filedialog.askopenfilename(
			title='Open audio file to transcribe',
			filetypes=[('Audio files', '*.mp3 *.wav'), ('mp3 file', '*.mp3'), ('wav file', '*.wav')],
		)
		if not chosen_file_path:
			return
		if not SR_AVAILABLE or Recognizer is None or AudioFile is None:
			messagebox.showerror(getattr(app, 'title_struct', '') + 'transcript', 'Speech recognition is not available.')
			return
		if not os.path.exists(chosen_file_path):
			messagebox.showerror(getattr(app, 'title_struct', '') + 'transcript', 'Selected file does not exist.')
			return

		base_name, extension = os.path.splitext(chosen_file_path)
		extension = extension.lower()
		if extension not in ('.mp3', '.wav'):
			messagebox.showerror(getattr(app, 'title_struct', '') + 'transcript', 'Unsupported file type. Choose an mp3 or wav file.')
			return

		try:
			file_size_bytes = os.path.getsize(chosen_file_path)
			if file_size_bytes >= 25 * 1024 * 1024:
				proceed_large = messagebox.askyesno(
					getattr(app, 'title_struct', '') + 'transcript',
					'This is a large file (~25MB+).\nTranscription may take a while.\nProceed?'
				)
				if not proceed_large:
					return
		except Exception:
			pass

		temp_wav_created = False
		wav_source_path = chosen_file_path

		try:
			if extension == '.mp3':
				if not PYDUB_AVAILABLE or AudioSegment is None:
					messagebox.showerror(
						getattr(app, 'title_struct', '') + 'transcript',
						'Transcribing mp3 requires pydub + ffmpeg. Please install them or use a wav file.'
					)
					return
				wav_source_path = base_name + '_converted.wav'
				AudioSegment.from_mp3(chosen_file_path).export(wav_source_path, format='wav')
				temp_wav_created = True
		except Exception as exc:
			messagebox.showerror(getattr(app, 'title_struct', '') + 'transcript', f'Failed preparing audio:\n{exc}')
			return

		cancel_event = Event()

		def task_transcribe():
			backoff_seconds = 1.0
			recognizer = Recognizer()
			language_code = stt_language_var.get().strip() or None

			if segment_audio_var.get():
				all_text_chunks = []
				current_offset = 0.0
				chunk_len = max(5, int(segment_seconds_var.get() or 90))
				for _ in range(0, 24 * 60 * 60 // chunk_len + 2):
					if cancel_event.is_set():
						break
					try:
						with AudioFile(wav_source_path) as audio_source:
							audio_data = recognizer.record(audio_source, offset=current_offset, duration=chunk_len)
						if not getattr(audio_data, 'frame_data', None) or len(audio_data.frame_data) < 64:
							break
						last_error_local = None
						for _attempt in range(3):
							try:
								if language_code:
									chunk_text = recognizer.recognize_google(audio_data, language=language_code)
								else:
									chunk_text = recognizer.recognize_google(audio_data)
								break
							except Exception as exc_local:
								last_error_local = exc_local
								time.sleep(backoff_seconds)
								backoff_seconds = min(8.0, backoff_seconds * 2.0)
						if last_error_local and 'chunk_text' not in locals():
							raise last_error_local
						cleaned_chunk = normalize_transcript_text(chunk_text)
						all_text_chunks.append(cleaned_chunk)
						current_offset += chunk_len
					except Exception:
						break
				final_text_value = '\n'.join(all_text_chunks)
				return final_text_value

			with AudioFile(wav_source_path) as audio_source:
				audio_blob = recognizer.record(audio_source)
			last_error = None
			for _attempt in range(3):
				if cancel_event.is_set():
					break
				try:
					if language_code:
						return normalize_transcript_text(recognizer.recognize_google(audio_blob, language=language_code))
					return normalize_transcript_text(recognizer.recognize_google(audio_blob))
				except Exception as exc_local:
					last_error = exc_local
					time.sleep(backoff_seconds)
					backoff_seconds = min(8.0, backoff_seconds * 2.0)
			if last_error:
				raise last_error
			return ''

		def after_transcribe(result_text, error_obj):
			try:
				if temp_wav_created and wav_source_path and os.path.exists(wav_source_path):
					os.remove(wav_source_path)
			except Exception:
				pass
			if cancel_event.is_set():
				messagebox.showinfo(getattr(app, 'title_struct', '') + 'transcript', 'Transcription canceled.')
				return
			if error_obj:
				messagebox.showerror(getattr(app, 'title_struct', '') + 'transcript', f'Failed to transcribe audio:\n{error_obj}')
				return
			render_text_result_window('File transcript', str(result_text or ''))

		run_async('Audio transcription', task_transcribe, on_done=after_transcribe, cancel_event=cancel_event)

	# --------------------- YouTube integration ---------------------
	def transcribe_youtube_from_ui(video_id_or_url: str, language_code: str, status_callback=None, disable_list=None):
		video_id_for_links = normalize_youtube_input(video_id_or_url)

		def task_fetch():
			backoff_seconds = 1.0
			last_error_fetch = None
			for _attempt in range(3):
				try:
					return fetch_youtube_captions(video_id_or_url, language_code)
				except Exception as exc:
					last_error_fetch = exc
					time.sleep(backoff_seconds)
					backoff_seconds = min(8.0, backoff_seconds * 2.0)
			raise last_error_fetch if last_error_fetch else RuntimeError('Unknown error')

		def after_fetch(result_items, error_obj):
			if error_obj:
				if callable(status_callback):
					status_callback('Failed to fetch transcript.', good=False)
				messagebox.showerror(getattr(app, 'title_struct', '') + 'transcript', str(error_obj))
				return
			if callable(status_callback):
				status_callback('Fetched successfully.', good=True)
			render_caption_result_window('YouTube transcript', result_items, source_video_id=video_id_for_links)

		run_async('YouTube', task_fetch, on_done=after_fetch, disable_widgets=(disable_list or []))

	# --------------------- selection window ---------------------
	def open_selection_window():
		selection_root = make_popup_window('Transcript', name='transcript_source')

		selection_title_label = Label(selection_root, text='Transcript source', font='arial 13 bold')
		selection_description_label = Label(
			selection_root,
			text='Paste a YouTube ID/URL (optional language) or transcribe a local audio file.',
			font='arial 9'
		)
		selection_title_label.pack(pady=(6, 2))
		selection_description_label.pack(pady=(0, 8))

		main_container_frame = Frame(selection_root)
		main_container_frame.pack(padx=10, pady=6)

		youtube_section_frame = Frame(main_container_frame, bd=1, relief='groove')
		youtube_section_frame.grid(row=0, column=0, sticky='nsew', padx=(0, 8))

		youtube_section_title_label = Label(youtube_section_frame, text='YouTube', font='arial 10 bold')
		youtube_section_title_label.grid(row=0, column=0, columnspan=2, sticky='w', pady=(4, 4))

		youtube_id_label = Label(youtube_section_frame, text='Video ID / URL:', font='arial 10')
		youtube_id_entry = Entry(youtube_section_frame, width=36)
		youtube_id_label.grid(row=1, column=0, sticky='w', pady=(0, 4))
		youtube_id_entry.grid(row=1, column=1, sticky='ew', padx=(6, 6), pady=(0, 4))
		youtube_section_frame.grid_columnconfigure(1, weight=1)

		youtube_language_label = Label(youtube_section_frame, text='Language (optional, e.g., en, es):', font='arial 9')
		youtube_language_entry = Entry(youtube_section_frame, textvariable=youtube_language_var, width=10)
		youtube_language_label.grid(row=2, column=0, sticky='w', pady=(0, 6))
		youtube_language_entry.grid(row=2, column=1, sticky='w', padx=(6, 0), pady=(0, 6))

		youtube_buttons_frame = Frame(youtube_section_frame)
		youtube_buttons_frame.grid(row=3, column=0, columnspan=2, sticky='w')
		youtube_fetch_button = Button(youtube_buttons_frame, text='Fetch', bd=1)
		youtube_clear_button = Button(youtube_buttons_frame, text='Clear', bd=1)
		youtube_paste_button = Button(youtube_buttons_frame, text='Paste', bd=1)
		youtube_fetch_button.grid(row=0, column=0, padx=(0, 8))
		youtube_clear_button.grid(row=0, column=1, padx=(0, 8))
		youtube_paste_button.grid(row=0, column=2)

		youtube_status_label = Label(youtube_section_frame, text='', font='arial 9')
		youtube_status_label.grid(row=4, column=0, columnspan=2, sticky='w', pady=(6, 8))

		def update_status(message_text, good=False):
			youtube_status_label.configure(fg=('dark green' if good else 'red'), text=message_text)

		def validate_youtube_entry(event=None):
			raw_value = youtube_id_entry.get().strip()
			normalized_value = normalize_youtube_input(raw_value)
			valid_value = looks_like_youtube_id(normalized_value)
			try:
				youtube_fetch_button.config(state='normal' if valid_value else 'disabled')
			except Exception:
				pass
			update_status('Ready' if valid_value else 'Enter a valid ID/URL', good=valid_value)

		def on_fetch_youtube():
			update_status('Fetching transcript...', good=False)
			transcribe_youtube_from_ui(
				youtube_id_entry.get().strip(),
				youtube_language_var.get().strip(),
				status_callback=update_status,
				disable_list=[youtube_fetch_button, youtube_clear_button, youtube_paste_button]
			)

		def on_clear_youtube():
			youtube_id_entry.delete(0, END)
			validate_youtube_entry()

		def on_paste_youtube():
			try:
				clipboard_text = selection_root.clipboard_get()
			except Exception:
				clipboard_text = ''
			if clipboard_text:
				youtube_id_entry.delete(0, END)
				youtube_id_entry.insert(0, clipboard_text.strip())
			validate_youtube_entry()

		def on_return_key(event=None):
			if youtube_fetch_button['state'] == 'normal':
				on_fetch_youtube()

		youtube_fetch_button.config(command=on_fetch_youtube, state='disabled')
		youtube_clear_button.config(command=on_clear_youtube)
		youtube_paste_button.config(command=on_paste_youtube)
		youtube_id_entry.bind('<KeyRelease>', validate_youtube_entry)
		youtube_id_entry.bind('<Return>', on_return_key)
		youtube_id_entry.focus_set()
		validate_youtube_entry()

		file_section_frame = Frame(main_container_frame, bd=1, relief='groove')
		file_section_frame.grid(row=0, column=1, sticky='nsew')

		file_section_title_label = Label(file_section_frame, text='Local audio file', font='arial 10 bold')
		file_hint_label = Label(file_section_frame, text='Choose an mp3 or wav file and transcribe it.', font='arial 10')
		file_browse_button = Button(file_section_frame, text='Browse...', bd=1, command=transcribe_local_audio_file)

		file_section_title_label.grid(row=0, column=0, sticky='w', pady=(4, 4))
		file_hint_label.grid(row=1, column=0, sticky='w', pady=(0, 6))
		file_browse_button.grid(row=2, column=0, sticky='w')

		options_section_frame = Frame(selection_root, bd=1, relief='groove')
		options_section_frame.pack(fill='x', padx=10, pady=(6, 0))

		redact_checkbutton = Checkbutton(options_section_frame, text='Redact emails and numbers', variable=redact_pii_var)
		show_timestamp_checkbutton = Checkbutton(options_section_frame, text='Show inline timestamps (TXT)', variable=show_inline_timestamp_var)
		redact_checkbutton.grid(row=0, column=0, sticky='w', padx=(6, 6))
		show_timestamp_checkbutton.grid(row=0, column=1, sticky='w', padx=(6, 6))

		stt_language_label = Label(options_section_frame, text='Local STT language (optional, e.g., en-US):', font='arial 9')
		stt_language_entry = Entry(options_section_frame, textvariable=stt_language_var, width=14)
		stt_language_label.grid(row=1, column=0, sticky='w', padx=(6, 6), pady=(6, 0))
		stt_language_entry.grid(row=1, column=1, sticky='w', padx=(6, 6), pady=(6, 0))

		segment_audio_checkbutton = Checkbutton(options_section_frame, text='Segment audio (seconds):', variable=segment_audio_var)
		segment_seconds_entry = Entry(options_section_frame, textvariable=segment_seconds_var, width=6)
		segment_audio_checkbutton.grid(row=1, column=2, sticky='w', padx=(24, 6), pady=(6, 0))
		segment_seconds_entry.grid(row=1, column=3, sticky='w', padx=(6, 6), pady=(6, 0))

		remove_bracketed_checkbutton = Checkbutton(options_section_frame, text='Remove bracketed tags [like this]', variable=remove_bracketed_var)
		collapse_whitespace_checkbutton = Checkbutton(options_section_frame, text='Collapse whitespace', variable=collapse_whitespace_var)
		remove_bracketed_checkbutton.grid(row=2, column=0, sticky='w', padx=(6, 6), pady=(6, 6))
		collapse_whitespace_checkbutton.grid(row=2, column=1, sticky='w', padx=(6, 6), pady=(6, 6))

		insert_mode_label = Label(options_section_frame, text='Default insert mode:', font='arial 9')
		insert_replace_radio = Radiobutton(options_section_frame, text='Replace selection', variable=insert_mode_var, value='replace')
		insert_cursor_radio = Radiobutton(options_section_frame, text='Insert at cursor', variable=insert_mode_var, value='insert')
		insert_append_radio = Radiobutton(options_section_frame, text='Append to end', variable=insert_mode_var, value='append')
		insert_mode_label.grid(row=3, column=0, sticky='w', pady=(6, 6), padx=(6, 6))
		insert_replace_radio.grid(row=3, column=1, padx=6, sticky='w')
		insert_cursor_radio.grid(row=3, column=2, padx=6, sticky='w')
		insert_append_radio.grid(row=3, column=3, padx=6, sticky='w')

		app.place_toolt([
			(youtube_fetch_button, 'Fetch transcript'),
			(youtube_clear_button, 'Clear input'),
			(youtube_paste_button, 'Paste from clipboard'),
			(file_browse_button, 'Browse local audio (mp3/wav)'),
			(show_timestamp_checkbutton, 'Toggle inline timestamps in TXT'),
			(segment_audio_checkbutton, 'Enable chunked transcription'),
		])

		privacy_note_label = Label(selection_root, text='Note: Transcripts may include personal data.', font='arial 8', fg='grey40')
		privacy_note_label.pack(pady=(6, 0))

		bottom_buttons_frame = Frame(selection_root)
		close_selection_button = Button(bottom_buttons_frame, text='Close', bd=1, command=selection_root.destroy)
		bottom_buttons_frame.pack(fill='x', padx=10, pady=(8, 10))
		close_selection_button.pack(side='right')

		try:
			main_container_frame.grid_columnconfigure(0, weight=1)
			main_container_frame.grid_columnconfigure(1, weight=1)
		except Exception:
			pass

		selection_root.bind('<Escape>', lambda event: (selection_root.destroy(), 'break'))

	if YT_API_AVAILABLE:
		open_selection_window()
	else:
		transcribe_local_audio_file()
