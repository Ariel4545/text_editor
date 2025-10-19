# file_template_generator_popup.py
from __future__ import annotations

import os
import platform
import shutil
import subprocess
import sys
import tkinter as tk
from collections import Counter, OrderedDict
from datetime import datetime
from pathlib import Path
from tkinter import ttk, filedialog, messagebox


def open_document_template_generator(app, event=None):
    '''
    Document Template Generator popup.

    Features:
    - Output format: PDF | WORD | JINJA
    - Subject presets: resume | personal info | computer stats | letter
    - Optional: Use docxtpl with a provided .docx template file
    - Edit placeholders before generating (dynamic context form)
    - Actions: Generate, Open Folder, Open File, Copy Path, Close
    '''
    text_widget = app.EgonTE

    # -- Constants for subjects and formats --
    SUBJECT_RESUME = 'resume'
    SUBJECT_PERSONAL_INFO = 'personal info'
    SUBJECT_COMPUTER_STATS = 'computer stats'
    SUBJECT_LETTER = 'letter'
    SUBJECT_MEETING_NOTES = 'meeting notes'
    SUBJECT_PROJECT_PROPOSAL = 'project proposal'
    SUBJECT_COVER_LETTER = 'cover letter'
    SUBJECT_CODE_REVIEW = 'code review'
    SUBJECTS = (
        SUBJECT_RESUME, SUBJECT_PERSONAL_INFO, SUBJECT_COMPUTER_STATS, SUBJECT_LETTER, SUBJECT_MEETING_NOTES,
        SUBJECT_PROJECT_PROPOSAL, SUBJECT_COVER_LETTER, SUBJECT_CODE_REVIEW
    )

    FORMAT_PDF = 'PDF'
    FORMAT_WORD = 'WORD'
    FORMAT_JINJA = 'JINJA'
    OUTPUT_FORMATS = (FORMAT_PDF, FORMAT_WORD, FORMAT_JINJA)


    # ---------------- Validation / import helpers ----------------
    def require_python_docx():
        '''
        Ensure the correct 'python-docx' package is installed (not the legacy 'docx').
        Returns: (Document, Pt)
        '''
        try:
            import docx  # provided by python-docx
            from docx import Document
            from docx.shared import Pt
        except ImportError as e:
            raise RuntimeError(
                "Word generation requires the 'python-docx' package.\n"
                "Install it using:\n"
                "  pip install python-docx"
            ) from e

        looks_legacy = False
        if not hasattr(docx, 'Document'):
            looks_legacy = True
        else:
            try:
                from docx import shared  # noqa: F401
                _ = Pt  # ensure available
            except Exception:
                looks_legacy = True

        if looks_legacy:
            raise RuntimeError(
                "It looks like the legacy 'docx' package is installed instead of 'python-docx'.\n"
                "Please run:\n"
                "  pip uninstall -y docx\n"
                "  pip install python-docx"
            )
        return Document, Pt

    def require_fpdf():
        try:
            from fpdf import FPDF
            return FPDF
        except ImportError as e:
            raise RuntimeError(
                "PDF generation requires the 'fpdf' package.\n"
                "Install it using:\n"
                "  pip install fpdf"
            ) from e

    def require_jinja2_template():
        try:
            from jinja2 import Template
            return Template
        except ImportError as e:
            raise RuntimeError(
                "Jinja rendering requires the 'jinja2' package.\n"
                "Install it using:\n"
                "  pip install jinja2"
            ) from e

    def require_docxtpl():
        try:
            from docxtpl import DocxTemplate
            return DocxTemplate
        except ImportError as e:
            raise RuntimeError(
                "Using a .docx template requires the 'docxtpl' package.\n"
                "Install it using:\n"
                "  pip install docxtpl"
            ) from e

    def require_psutil():
        """Safely import psutil, returning None if not available."""
        try:
            import psutil
            return psutil
        except ImportError:
            return None

    def require_gputil():
        """Safely import GPUtil, returning None if not available."""
        try:
            import GPUtil
            return GPUtil
        except ImportError:
            return None


    # ---------------- Subject Context Definitions ----------------
    def _get_resume_context():
        return 'Resume', OrderedDict([
            ('full_name', '<FULL_NAME>'),
            ('job_title', '<JOB_TITLE>'),
            ('contact_info', '<EMAIL> | <PHONE> | <LINKEDIN_URL>'),
            ('summary_heading', 'Summary'),
            ('summary_body', '<A few sentences summarizing your skills and experience.>'),
            ('skills_heading', 'Skills'),
            ('skills_body', 'Languages: Python, SQL, JavaScript\nFrameworks: Django, React, Node.js\nTools: Git, Docker, AWS'),
            ('experience_heading', 'Experience'),
            ('experience_body', 'Job Title; Company; Date Range\n• Description point 1\n• Description point 2\n---\nAnother Job Title; Another Company; Date Range\n• Description point 1'),
            ('education_heading', 'Education'),
            ('education_body', 'Degree or Certificate; School; Graduation Year'),
        ])

    def _get_personal_info_context():
        return 'Personal Information', OrderedDict([
            ('full_name', '<FULL_NAME>'),
            ('address', '<ADDRESS_LINE>'),
            ('city', '<CITY>'),
            ('country', '<COUNTRY>'),
            ('email', '<EMAIL>'),
            ('phone', '<PHONE>'),
            ('dob', '<YYYY-MM-DD>'),
        ])

    def _get_computer_stats_context():
        try:
            python_version = sys.version.split()[0]
        except IndexError:
            python_version = '<PYTHON_VERSION>'

        ram_gb = '<RAM_GB> GB'
        psutil = require_psutil()
        if psutil:
            try:
                ram_total_bytes = psutil.virtual_memory().total
                ram_gb = f'{ram_total_bytes / (1024 ** 3):.1f} GB'
            except Exception:
                pass  # Fallback to placeholder

        storage_gb = '<STORAGE_TOTAL_GB> GB'
        try:
            total, _, _ = shutil.disk_usage('/')
            storage_gb = f'{total / (1024 ** 3):.1f} GB'
        except Exception:
            pass  # Fallback to placeholder

        gpu_model = '<GPU_MODEL>'
        gputil = require_gputil()
        if gputil:
            try:
                gpus = gputil.getGPUs()
                if gpus:
                    gpu_model = gpus[0].name
            except Exception:
                pass  # Fallback to placeholder

        return 'Computer Stats', OrderedDict([
            ('os', f'{platform.system()} {platform.release()}' or '<OS_NAME_VERSION>'),
            ('cpu', platform.processor() or '<CPU_MODEL>'),
            ('ram', ram_gb),
            ('storage', storage_gb),
            ('gpu', gpu_model),
            ('python', python_version),
        ])

    def _get_letter_context():
        return 'Letter', OrderedDict([
            ('sender_name', '<YOUR_NAME>'),
            ('date', datetime.now().strftime('%Y-%m-%d')),
            ('recipient_name', '<RECIPIENT_NAME>'),
            ('recipient_company', '<RECIPIENT_COMPANY>'),
            ('greeting', 'Dear <RECIPIENT_NAME>,'),
            ('body_intro', '<Opening paragraph...>'),
            ('body_main', '<Main paragraphs...>'),
            ('body_conclusion', '<Closing paragraph...>'),
            ('closing', 'Sincerely,'),
        ])

    def _get_cover_letter_context():
        return 'Cover Letter', OrderedDict([
            ('your_name', '<YOUR_NAME>'),
            ('your_address', '<YOUR_STREET_ADDRESS, CITY, ZIP>'),
            ('your_contact_info', '<YOUR_EMAIL> | <YOUR_PHONE>'),
            ('date', datetime.now().strftime('%Y-%m-%d')),
            ('hiring_manager_heading', 'Recipient'),
            ('hiring_manager_name', '<HIRING_MANAGER_NAME (if known)>'),
            ('hiring_manager_title', '<HIRING_MANAGER_TITLE>'),
            ('company_name', '<COMPANY_NAME>'),
            ('company_address', '<COMPANY_STREET_ADDRESS, CITY, ZIP>'),
            ('body_heading', 'Letter Body'),
            ('greeting', 'Dear <Mr./Ms./Mx. LAST_NAME>,'),
            ('body_intro', '<State the position you are applying for and where you saw the opening.>'),
            ('body_main', '<Highlight your most relevant skills and experiences that match the job description.>'),
            ('body_conclusion', '<Express your enthusiasm for the role and the company. Mention why you are a good fit for their culture.>'),
            ('call_to_action', 'I am eager to learn more about this opportunity and discuss how my skills can benefit your team. I am available for an interview at your earliest convenience.'),
            ('closing', 'Sincerely,'),
        ])

    def _get_meeting_notes_context():
        return 'Meeting Notes', OrderedDict([
            ('meeting_title', '<MEETING_TITLE>'),
            ('date', datetime.now().strftime('%Y-%m-%d')),
            ('attendees', '<ATTENDEE_1>, <ATTENDEE_2>'),
            ('agenda_heading', 'Agenda'),
            ('agenda_items', '• <TOPIC_1>\n• <TOPIC_2>'),
            ('notes_heading', 'Notes'),
            ('notes_body', '<NOTES_TAKEN_DURING_THE_MEETING>'),
            ('actions_heading', 'Action Items'),
            ('action_items', '• [ ] <ACTION_ITEM_1> (Owner: <NAME>)'),
        ])

    def _get_project_proposal_context():
        return 'Project Proposal', OrderedDict([
            ('project_title', '<PROJECT_TITLE>'),
            ('date', datetime.now().strftime('%Y-%m-%d')),
            ('client_name', '<CLIENT_NAME>'),
            ('your_company_name', '<YOUR_COMPANY_NAME>'),
            ('introduction_heading', 'Introduction'),
            ('introduction_body', '<Briefly introduce the project and its purpose.>'),
            ('problem_statement_heading', 'Problem Statement'),
            ('problem_statement_body', '<Describe the problem this project will solve.>'),
            ('solution_heading', 'Proposed Solution'),
            ('solution_body', '<Detail your proposed solution.>'),
            ('deliverables_heading', 'Deliverables'),
            ('deliverables_items', '• <DELIVERABLE_1>\n• <DELIVERABLE_2>'),
            ('timeline_heading', 'Timeline'),
            ('timeline_body', '<Provide an estimated timeline for the project.>'),
            ('budget_heading', 'Budget'),
            ('budget_items', 'Phase 1: Discovery, 10 hours, 1500\nPhase 2: Development, 40 hours, 6000'),
            ('budget_summary', '<Optional summary of costs or payment terms.>'),
            ('conclusion_heading', 'Conclusion'),
            ('conclusion_body', '<Summarize the proposal and next steps.>'),
        ])

    def _get_code_review_context():
        return 'Code Review Checklist', OrderedDict([
            ('project_name', '<PROJECT_NAME>'),
            ('feature_name', '<FEATURE_OR_CHANGE_NAME>'),
            ('author', '<AUTHOR_NAME>'),
            ('reviewer', '<YOUR_NAME>'),
            ('date', datetime.now().strftime('%Y-%m-%d')),
            ('summary_heading', 'Summary of Changes'),
            ('summary_body', '<Briefly describe the purpose of the code change.>'),
            ('correctness_heading', 'Correctness'),
            ('correctness_items', '• [ ] Does the code do what it says it does?\n• [ ] Are edge cases handled correctly?\n• [ ] Is the logic sound?'),
            ('readability_heading', 'Readability & Style'),
            ('readability_items', '• [ ] Is the code easy to understand?\n• [ ] Does it follow the project\'s style guide?\n• [ ] Are variable and function names clear and descriptive?'),
            ('testing_heading', 'Testing'),
            ('testing_items', '• [ ] Are there sufficient tests for the new code?\n• [ ] Do all tests pass?\n• [ ] Are the tests well-written?'),
            ('comments_heading', 'Overall Comments'),
            ('comments_body', '<General feedback or summary of the review.>'),
        ])

    _SUBJECT_CONTEXT_GENERATORS = {
        SUBJECT_RESUME: _get_resume_context,
        SUBJECT_PERSONAL_INFO: _get_personal_info_context,
        SUBJECT_COMPUTER_STATS: _get_computer_stats_context,
        SUBJECT_LETTER: _get_letter_context,
        SUBJECT_COVER_LETTER: _get_cover_letter_context,
        SUBJECT_MEETING_NOTES: _get_meeting_notes_context,
        SUBJECT_PROJECT_PROPOSAL: _get_project_proposal_context,
        SUBJECT_CODE_REVIEW: _get_code_review_context,
    }

    def get_default_context_for_subject(subject_key: str) -> tuple[str, dict]:
        generator = _SUBJECT_CONTEXT_GENERATORS.get(subject_key)
        if generator:
            return generator()
        return subject_key.title(), {'note': '<PLACEHOLDER>'}


    # ---------------- Content / generation helpers ----------------
    def ensure_output_dir() -> Path:
        out_path = Path('outputs')
        out_path.mkdir(parents=True, exist_ok=True)
        return out_path

    def build_output_filename(prefix: str, ext: str) -> Path:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_prefix = str(prefix).replace(' ', '_').lower()
        return ensure_output_dir() / f'{safe_prefix}_{timestamp}.{ext}'

    def _add_bullet_points(doc_or_pdf, text: str, is_pdf: bool):
        lines = text.split('\n')
        for line in lines:
            cleaned_line = line.lstrip('•*- ').strip()
            if not cleaned_line:
                continue
            if is_pdf:
                doc_or_pdf.multi_cell(0, 5, txt=f'  •  {cleaned_line}')
            else:
                doc_or_pdf.add_paragraph(cleaned_line, style='List Bullet')

    def _generate_resume_pdf(title: str, context: dict) -> str:
        FPDF = require_fpdf()
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_left_margin(20)
        pdf.set_right_margin(20)

        pdf.set_font('Arial', 'B', 18)
        pdf.cell(0, 8, txt=context.get('full_name', ''), ln=True, align='C')
        pdf.set_font('Arial', '', 12)
        pdf.cell(0, 6, txt=context.get('job_title', ''), ln=True, align='C')
        pdf.cell(0, 6, txt=context.get('contact_info', ''), ln=True, align='C')
        pdf.ln(10)

        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, txt=context.get('summary_heading', 'Summary'), ln=True, border='B')
        pdf.ln(4)
        pdf.set_font('Arial', '', 11)
        pdf.multi_cell(0, 6, txt=context.get('summary_body', ''))
        pdf.ln(4)

        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, txt=context.get('skills_heading', 'Skills'), ln=True, border='B')
        pdf.ln(4)
        pdf.set_font('Arial', '', 11)
        for line in context.get('skills_body', '').split('\n'):
            parts = line.split(':', 1)
            if len(parts) == 2:
                pdf.set_font('Arial', 'B', 11)
                pdf.cell(pdf.get_string_width(parts[0] + ': '), 6, txt=parts[0] + ':')
                pdf.set_font('Arial', '', 11)
                pdf.multi_cell(0, 6, txt=parts[1].strip())
            else:
                pdf.multi_cell(0, 6, txt=line)
        pdf.ln(4)

        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, txt=context.get('experience_heading', 'Experience'), ln=True, border='B')
        pdf.ln(4)
        for entry in context.get('experience_entries', []):
            pdf.set_font('Arial', 'B', 12)
            pdf.multi_cell(0, 6, txt=entry.get('title', ''))
            pdf.set_font('Arial', 'I', 11)
            pdf.multi_cell(0, 6, txt=entry.get('company_info', ''))
            pdf.ln(2)
            pdf.set_font('Arial', '', 11)
            _add_bullet_points(pdf, entry.get('description', ''), is_pdf=True)
            pdf.ln(4)

        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, txt=context.get('education_heading', 'Education'), ln=True, border='B')
        pdf.ln(4)
        for entry in context.get('education_entries', []):
            pdf.set_font('Arial', 'B', 12)
            pdf.multi_cell(0, 6, txt=entry.get('degree', ''))
            pdf.set_font('Arial', 'I', 11)
            pdf.multi_cell(0, 6, txt=entry.get('school_info', ''))
            pdf.ln(4)

        path = build_output_filename(title, 'pdf')
        pdf.output(str(path))
        return str(path)

    def _generate_letter_pdf(title: str, context: dict) -> str:
        FPDF = require_fpdf()
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font('Arial', size=12)

        if context.get('sender_name'):
            pdf.multi_cell(0, 5, txt=context['sender_name'], align='R')
        if context.get('date'):
            pdf.multi_cell(0, 5, txt=context['date'], align='R')
        pdf.ln(10)

        if context.get('recipient_name'):
            pdf.cell(0, 5, txt=context['recipient_name'], ln=True)
        if context.get('recipient_company'):
            pdf.cell(0, 5, txt=context['recipient_company'], ln=True)
        pdf.ln(10)

        for key in ['greeting', 'body_intro', 'body_main', 'body_conclusion']:
            if context.get(key):
                pdf.multi_cell(0, 5, txt=context[key])
                pdf.ln(5)

        if context.get('closing'):
            pdf.cell(0, 5, txt=context['closing'], ln=True)
        if context.get('sender_name'):
            pdf.cell(0, 5, txt=context['sender_name'], ln=True)

        path = build_output_filename(title, 'pdf')
        pdf.output(str(path))
        return str(path)

    def _generate_cover_letter_pdf(title: str, context: dict) -> str:
        FPDF = require_fpdf()
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font('Arial', size=12)

        for key in ['your_name', 'your_address', 'your_contact_info', 'date']:
            pdf.multi_cell(0, 5, txt=context.get(key, ''))
        pdf.ln(10)

        for key in ['hiring_manager_name', 'hiring_manager_title', 'company_name', 'company_address']:
            pdf.multi_cell(0, 5, txt=context.get(key, ''))
        pdf.ln(10)

        for key in ['greeting', 'body_intro', 'body_main', 'body_conclusion', 'call_to_action']:
            pdf.multi_cell(0, 5, txt=context.get(key, ''))
            pdf.ln(5)

        pdf.multi_cell(0, 5, txt=context.get('closing', ''))
        pdf.ln(5)
        pdf.multi_cell(0, 5, txt=context.get('your_name', ''))

        path = build_output_filename(title, 'pdf')
        pdf.output(str(path))
        return str(path)

    def _generate_meeting_notes_pdf(title: str, context: dict) -> str:
        FPDF = require_fpdf()
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        pdf.set_font('Arial', 'B', 18)
        pdf.cell(0, 8, txt=context.get('meeting_title', title), ln=True, align='C')
        pdf.ln(5)

        pdf.set_font('Arial', '', 12)
        pdf.cell(0, 6, txt=f"Date: {context.get('date', '')}", ln=True)
        pdf.cell(0, 6, txt=f"Attendees: {context.get('attendees', '')}", ln=True)
        pdf.ln(10)

        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, txt=context.get('agenda_heading', 'Agenda'), ln=True, border='B')
        pdf.ln(4)
        pdf.set_font('Arial', '', 11)
        _add_bullet_points(pdf, context.get('agenda_items', ''), is_pdf=True)
        pdf.ln(4)

        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, txt=context.get('notes_heading', 'Notes'), ln=True, border='B')
        pdf.ln(4)
        pdf.set_font('Arial', '', 11)
        pdf.multi_cell(0, 6, txt=context.get('notes_body', ''))
        pdf.ln(4)

        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, txt=context.get('actions_heading', 'Action Items'), ln=True, border='B')
        pdf.ln(4)
        pdf.set_font('Arial', '', 11)
        _add_bullet_points(pdf, context.get('action_items', ''), is_pdf=True)

        path = build_output_filename(title, 'pdf')
        pdf.output(str(path))
        return str(path)

    def _generate_project_proposal_pdf(title: str, context: dict) -> str:
        FPDF = require_fpdf()
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        pdf.set_font('Arial', 'B', 18)
        pdf.cell(0, 8, txt=context.get('project_title', title), ln=True, align='C')
        pdf.ln(5)

        pdf.set_font('Arial', '', 12)
        pdf.cell(0, 6, txt=f"Date: {context.get('date', '')}", ln=True)
        pdf.cell(0, 6, txt=f"Client: {context.get('client_name', '')}", ln=True)
        pdf.cell(0, 6, txt=f"From: {context.get('your_company_name', '')}", ln=True)
        pdf.ln(10)

        sections = ['introduction', 'problem_statement', 'solution', 'deliverables', 'timeline', 'budget', 'conclusion']
        for section in sections:
            pdf.set_font('Arial', 'B', 14)
            pdf.cell(0, 10, txt=context.get(f'{section}_heading', section.title()), ln=True, border='B')
            pdf.ln(4)
            pdf.set_font('Arial', '', 11)

            if section == 'deliverables':
                _add_bullet_points(pdf, context.get('deliverables_items', ''), is_pdf=True)
            elif section == 'budget':
                budget_entries = context.get('budget_entries', [])
                if budget_entries:
                    col_widths = (pdf.w - pdf.l_margin - pdf.r_margin) * 0.4, (pdf.w - pdf.l_margin - pdf.r_margin) * 0.4, (pdf.w - pdf.l_margin - pdf.r_margin) * 0.2
                    pdf.set_font('Arial', 'B', 11)
                    pdf.cell(col_widths[0], 8, 'Item', border=1)
                    pdf.cell(col_widths[1], 8, 'Description', border=1)
                    pdf.cell(col_widths[2], 8, 'Cost', border=1, ln=1)
                    pdf.set_font('Arial', '', 11)
                    for item in budget_entries:
                        pdf.cell(col_widths[0], 8, item.get('name', ''), border=1)
                        pdf.cell(col_widths[1], 8, item.get('desc', ''), border=1)
                        pdf.cell(col_widths[2], 8, str(item.get('cost', '0')), border=1, ln=1)
                    pdf.set_font('Arial', 'B', 11)
                    pdf.cell(col_widths[0] + col_widths[1], 8, 'Total Estimated Cost', border=1)
                    pdf.cell(col_widths[2], 8, f"{context.get('budget_total', 0.0):.2f}", border=1, ln=1)
                pdf.ln(4)
                pdf.set_font('Arial', '', 11)
                pdf.multi_cell(0, 6, txt=context.get('budget_summary', ''))
            else:
                pdf.multi_cell(0, 6, txt=context.get(f'{section}_body', ''))
            pdf.ln(4)

        path = build_output_filename(title, 'pdf')
        pdf.output(str(path))
        return str(path)

    def _generate_code_review_pdf(title: str, context: dict) -> str:
        FPDF = require_fpdf()
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        pdf.set_font('Arial', 'B', 18)
        pdf.cell(0, 8, txt=f"Code Review: {context.get('feature_name', title)}", ln=True, align='C')
        pdf.ln(5)

        pdf.set_font('Arial', '', 12)
        pdf.cell(0, 6, txt=f"Project: {context.get('project_name', '')}", ln=True)
        pdf.cell(0, 6, txt=f"Author: {context.get('author', '')}", ln=True)
        pdf.cell(0, 6, txt=f"Reviewer: {context.get('reviewer', '')}", ln=True)
        pdf.cell(0, 6, txt=f"Date: {context.get('date', '')}", ln=True)
        pdf.ln(10)

        # Declarative structure
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, txt=context.get('summary_heading', ''), ln=True, border='B')
        pdf.ln(4)
        pdf.set_font('Arial', '', 11)
        pdf.multi_cell(0, 6, txt=context.get('summary_body', ''))
        pdf.ln(4)

        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, txt=context.get('correctness_heading', ''), ln=True, border='B')
        pdf.ln(4)
        pdf.set_font('Arial', '', 11)
        _add_bullet_points(pdf, context.get('correctness_items', ''), is_pdf=True)
        pdf.ln(4)

        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, txt=context.get('readability_heading', ''), ln=True, border='B')
        pdf.ln(4)
        pdf.set_font('Arial', '', 11)
        _add_bullet_points(pdf, context.get('readability_items', ''), is_pdf=True)
        pdf.ln(4)

        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, txt=context.get('testing_heading', ''), ln=True, border='B')
        pdf.ln(4)
        pdf.set_font('Arial', '', 11)
        _add_bullet_points(pdf, context.get('testing_items', ''), is_pdf=True)
        pdf.ln(4)

        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, txt=context.get('comments_heading', ''), ln=True, border='B')
        pdf.ln(4)
        pdf.set_font('Arial', '', 11)
        pdf.multi_cell(0, 6, txt=context.get('comments_body', ''))
        pdf.ln(4)

        path = build_output_filename(title, 'pdf')
        pdf.output(str(path))
        return str(path)

    def _generate_generic_pdf(title: str, context: dict) -> str:
        FPDF = require_fpdf()
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font('Arial', 'B', 16)
        pdf.cell(0, 10, txt=title, ln=True)
        pdf.ln(4)
        pdf.set_font('Arial', size=12)
        for key, value in context.items():
            pdf.multi_cell(0, 8, f'{key.replace("_", " ").title()}: {value}')
        path = build_output_filename(title, 'pdf')
        pdf.output(str(path))
        return str(path)

    def _generate_meeting_notes_word(title: str, context: dict) -> str:
        Document, Pt = require_python_docx()
        doc = Document()
        doc.add_heading(context.get('meeting_title', title), level=1)
        doc.add_paragraph(f"Date: {context.get('date', '')}")
        doc.add_paragraph(f"Attendees: {context.get('attendees', '')}")

        doc.add_heading(context.get('agenda_heading', 'Agenda'), level=2)
        _add_bullet_points(doc, context.get('agenda_items', ''), is_pdf=False)

        doc.add_heading(context.get('notes_heading', 'Notes'), level=2)
        doc.add_paragraph(context.get('notes_body', ''))

        doc.add_heading(context.get('actions_heading', 'Action Items'), level=2)
        _add_bullet_points(doc, context.get('action_items', ''), is_pdf=False)

        path = build_output_filename(title, 'docx')
        doc.save(str(path))
        return str(path)

    def _generate_letter_word(title: str, context: dict) -> str:
        Document, Pt = require_python_docx()
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            WD_ALIGN_PARAGRAPH = None

        doc = Document()

        if context.get('sender_name'):
            p = doc.add_paragraph(context['sender_name'])
            if WD_ALIGN_PARAGRAPH: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if context.get('date'):
            p = doc.add_paragraph(context['date'])
            if WD_ALIGN_PARAGRAPH: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph()  # Spacer

        if context.get('recipient_name'):
            doc.add_paragraph(context['recipient_name'])
        if context.get('recipient_company'):
            doc.add_paragraph(context['recipient_company'])
        doc.add_paragraph()

        for key in ['greeting', 'body_intro', 'body_main', 'body_conclusion']:
            if context.get(key):
                doc.add_paragraph(context[key])
                if 'body' in key: doc.add_paragraph()

        if context.get('closing'):
            doc.add_paragraph(context['closing'])
        if context.get('sender_name'):
            doc.add_paragraph(context['sender_name'])

        path = build_output_filename(title, 'docx')
        doc.save(str(path))
        return str(path)

    def _generate_cover_letter_word(title: str, context: dict) -> str:
        Document, Pt = require_python_docx()
        doc = Document()

        for key in ['your_name', 'your_address', 'your_contact_info', 'date']:
            doc.add_paragraph(context.get(key, ''))
        doc.add_paragraph()

        for key in ['hiring_manager_name', 'hiring_manager_title', 'company_name', 'company_address']:
            doc.add_paragraph(context.get(key, ''))
        doc.add_paragraph()

        for key in ['greeting', 'body_intro', 'body_main', 'body_conclusion', 'call_to_action']:
            doc.add_paragraph(context.get(key, ''))
            doc.add_paragraph()

        doc.add_paragraph(context.get('closing', ''))
        doc.add_paragraph()
        doc.add_paragraph(context.get('your_name', ''))

        path = build_output_filename(title, 'docx')
        doc.save(str(path))
        return str(path)

    def _generate_resume_word(title: str, context: dict) -> str:
        Document, Pt = require_python_docx()
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Inches
        except ImportError:
            WD_ALIGN_PARAGRAPH = None
            Inches = None

        doc = Document()
        if Inches:
            doc.sections[0].left_margin = Inches(0.75)
            doc.sections[0].right_margin = Inches(0.75)

        p = doc.add_paragraph()
        p.add_run(context.get('full_name', '')).bold = True
        if WD_ALIGN_PARAGRAPH: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph(context.get('job_title', ''))
        if WD_ALIGN_PARAGRAPH: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph(context.get('contact_info', ''))
        if WD_ALIGN_PARAGRAPH: p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading(context.get('summary_heading', 'Summary'), level=2)
        doc.add_paragraph(context.get('summary_body', ''))

        doc.add_heading(context.get('skills_heading', 'Skills'), level=2)
        p = doc.add_paragraph()
        for line in context.get('skills_body', '').split('\n'):
            parts = line.split(':', 1)
            if len(parts) == 2:
                p.add_run(f'{parts[0]}: ').bold = True
                p.add_run(f'{parts[1].strip()}\n')
            else:
                p.add_run(line + '\n')

        doc.add_heading(context.get('experience_heading', 'Experience'), level=2)
        for entry in context.get('experience_entries', []):
            p = doc.add_paragraph()
            p.add_run(entry.get('title', '')).bold = True
            p = doc.add_paragraph()
            p.add_run(entry.get('company_info', '')).italic = True
            _add_bullet_points(doc, entry.get('description', ''), is_pdf=False)

        doc.add_heading(context.get('education_heading', 'Education'), level=2)
        for entry in context.get('education_entries', []):
            p = doc.add_paragraph()
            p.add_run(entry.get('degree', '')).bold = True
            p = doc.add_paragraph()
            p.add_run(entry.get('school_info', '')).italic = True

        path = build_output_filename(title, 'docx')
        doc.save(str(path))
        return str(path)

    def _generate_project_proposal_word(title: str, context: dict) -> str:
        Document, Pt = require_python_docx()
        doc = Document()

        doc.add_heading(context.get('project_title', title), level=1)
        doc.add_paragraph(f"Date: {context.get('date', '')}")
        doc.add_paragraph(f"Client: {context.get('client_name', '')}")
        doc.add_paragraph(f"From: {context.get('your_company_name', '')}")

        sections = ['introduction', 'problem_statement', 'solution', 'deliverables', 'timeline', 'budget', 'conclusion']
        for section in sections:
            doc.add_heading(context.get(f'{section}_heading', section.title()), level=2)
            if section == 'deliverables':
                _add_bullet_points(doc, context.get('deliverables_items', ''), is_pdf=False)
            elif section == 'budget':
                budget_entries = context.get('budget_entries', [])
                if budget_entries:
                    table = doc.add_table(rows=1, cols=3)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Item'
                    hdr_cells[1].text = 'Description'
                    hdr_cells[2].text = 'Cost'
                    for item in budget_entries:
                        row_cells = table.add_row().cells
                        row_cells[0].text = item.get('name', '')
                        row_cells[1].text = item.get('desc', '')
                        row_cells[2].text = str(item.get('cost', '0'))

                    row_cells = table.add_row().cells
                    row_cells[0].merge(row_cells[1])
                    row_cells[0].text = 'Total Estimated Cost'
                    row_cells[0].paragraphs[0].runs[0].bold = True
                    row_cells[2].text = f"{context.get('budget_total', 0.0):.2f}"
                    row_cells[2].paragraphs[0].runs[0].bold = True
                doc.add_paragraph(context.get('budget_summary', ''))
            else:
                doc.add_paragraph(context.get(f'{section}_body', ''))

        path = build_output_filename(title, 'docx')
        doc.save(str(path))
        return str(path)

    def _generate_code_review_word(title: str, context: dict) -> str:
        Document, Pt = require_python_docx()
        doc = Document()
        doc.add_heading(f"Code Review: {context.get('feature_name', title)}", level=1)

        doc.add_paragraph(f"Project: {context.get('project_name', '')}")
        doc.add_paragraph(f"Author: {context.get('author', '')}")
        doc.add_paragraph(f"Reviewer: {context.get('reviewer', '')}")
        doc.add_paragraph(f"Date: {context.get('date', '')}")

        doc.add_heading(context.get('summary_heading', ''), level=2)
        doc.add_paragraph(context.get('summary_body', ''))

        doc.add_heading(context.get('correctness_heading', ''), level=2)
        _add_bullet_points(doc, context.get('correctness_items', ''), is_pdf=False)

        doc.add_heading(context.get('readability_heading', ''), level=2)
        _add_bullet_points(doc, context.get('readability_items', ''), is_pdf=False)

        doc.add_heading(context.get('testing_heading', ''), level=2)
        _add_bullet_points(doc, context.get('testing_items', ''), is_pdf=False)

        doc.add_heading(context.get('comments_heading', ''), level=2)
        doc.add_paragraph(context.get('comments_body', ''))

        path = build_output_filename(title, 'docx')
        doc.save(str(path))
        return str(path)

    def _generate_generic_word(title: str, context: dict) -> str:
        Document, Pt = require_python_docx()
        doc = Document()
        doc.add_heading(title, level=1)
        for key, value in context.items():
            p = doc.add_paragraph()
            p.add_run(f'{key.replace("_", " ").title()}: ').bold = True
            p.add_run(str(value))
        path = build_output_filename(title, 'docx')
        doc.save(str(path))
        return str(path)

    _JINJA_TEMPLATE_GENERIC = (
        '{{ title }}\n'
        '{% for _ in range(title|length) %}={% endfor %}\n\n'
        "{% for k, v in context.items() %}{{ k.replace('_',' ') | title }}: {{ v }}\n{% endfor %}"
    )

    _JINJA_TEMPLATE_RESUME = """
# {{ context.get('full_name', 'Resume') | upper }}

**{{ context.get('job_title', '') }}**
*{{ context.get('contact_info', '') }}*

## {{ context.get('summary_heading', 'Summary') }}
{{ context.get('summary_body', '') }}

## {{ context.get('skills_heading', 'Skills') }}
{{ context.get('skills_body', '') }}

## {{ context.get('experience_heading', 'Experience') }}
{% for job in context.experience_entries %}
### {{ job.title }}
*{{ job.company_info }}*
{{ job.description }}
{% endfor %}

## {{ context.get('education_heading', 'Education') }}
{% for edu in context.education_entries %}
### {{ edu.degree }}
*{{ edu.school_info }}*
{% endfor %}
    """

    _JINJA_TEMPLATE_LETTER = """
{{ context.get('sender_name', '') }}
{{ context.get('date', '') }}

{{ context.get('recipient_name', '') }}
{{ context.get('recipient_company', '') }}

{{ context.get('greeting', '') }}

{{ context.get('body_intro', '') }}

{{ context.get('body_main', '') }}

{{ context.get('body_conclusion', '') }}

{{ context.get('closing', '') }}
{{ context.get('sender_name', '') }}
    """

    _JINJA_TEMPLATE_COVER_LETTER = """
{{ context.get('your_name', '') }}
{{ context.get('your_address', '') }}
{{ context.get('your_contact_info', '') }}
{{ context.get('date', '') }}

{{ context.get('hiring_manager_name', '') }}
{{ context.get('hiring_manager_title', '') }}
{{ context.get('company_name', '') }}
{{ context.get('company_address', '') }}

{{ context.get('greeting', '') }}

{{ context.get('body_intro', '') }}

{{ context.get('body_main', '') }}

{{ context.get('body_conclusion', '') }}

{{ context.get('call_to_action', '') }}

{{ context.get('closing', '') }}

{{ context.get('your_name', '') }}
    """

    _JINJA_TEMPLATE_MEETING_NOTES = """
# {{ context.get('meeting_title', 'Meeting Notes') }}

**Date:** {{ context.get('date', '') }}
**Attendees:** {{ context.get('attendees', '') }}

## {{ context.get('agenda_heading', 'Agenda') }}
{{ context.get('agenda_items', '') }}

## {{ context.get('notes_heading', 'Notes') }}
{{ context.get('notes_body', '') }}

## {{ context.get('actions_heading', 'Action Items') }}
{{ context.get('action_items', '') }}
    """

    _JINJA_TEMPLATE_PROJECT_PROPOSAL = """
# {{ context.get('project_title', 'Project Proposal') }}

**Date:** {{ context.get('date', '') }}
**Client:** {{ context.get('client_name', '') }}
**From:** {{ context.get('your_company_name', '') }}

## {{ context.get('introduction_heading', 'Introduction') }}
{{ context.get('introduction_body', '') }}

## {{ context.get('problem_statement_heading', 'Problem Statement') }}
{{ context.get('problem_statement_body', '') }}

## {{ context.get('solution_heading', 'Proposed Solution') }}
{{ context.get('solution_body', '') }}

## {{ context.get('deliverables_heading', 'Deliverables') }}
{{ context.get('deliverables_items', '') }}

## {{ context.get('timeline_heading', 'Timeline') }}
{{ context.get('timeline_body', '') }}

## {{ context.get('budget_heading', 'Budget') }}
{{ context.get('budget_items', '') }}

*{{ context.get('budget_summary', '') }}*

## {{ context.get('conclusion_heading', 'Conclusion') }}
{{ context.get('conclusion_body', '') }}
    """

    _JINJA_TEMPLATE_CODE_REVIEW = """
# Code Review: {{ context.get('feature_name', '') }} for {{ context.get('project_name', '') }}

**Author:** {{ context.get('author', '') }}
**Reviewer:** {{ context.get('reviewer', '') }}
**Date:** {{ context.get('date', '') }}

## {{ context.get('summary_heading', 'Summary of Changes') }}
{{ context.get('summary_body', '') }}

## {{ context.get('correctness_heading', 'Correctness') }}
{{ context.get('correctness_items', '') }}

## {{ context.get('readability_heading', 'Readability & Style') }}
{{ context.get('readability_items', '') }}

## {{ context.get('testing_heading', 'Testing') }}
{{ context.get('testing_items', '') }}

## {{ context.get('comments_heading', 'Overall Comments') }}
{{ context.get('comments_body', '') }}
    """

    def _render_jinja_template(title: str, context: dict, template_text: str, ext: str = 'txt') -> str:
        Template = require_jinja2_template()
        rendered = Template(template_text, trim_blocks=True, lstrip_blocks=True).render(title=title, context=context)
        path = build_output_filename(title, ext)
        try:
            Path(path).write_text(rendered, encoding='utf-8')
        except Exception as e:
            raise IOError(f"Failed to write to file: {path}") from e
        return str(path)

    def _generate_generic_jinja(title: str, context: dict) -> str:
        return _render_jinja_template(title, context, _JINJA_TEMPLATE_GENERIC)

    def _generate_resume_jinja(title: str, context: dict) -> str:
        return _render_jinja_template(title, context, _JINJA_TEMPLATE_RESUME)

    def _generate_letter_jinja(title: str, context: dict) -> str:
        return _render_jinja_template(title, context, _JINJA_TEMPLATE_LETTER)

    def _generate_cover_letter_jinja(title: str, context: dict) -> str:
        return _render_jinja_template(title, context, _JINJA_TEMPLATE_COVER_LETTER)

    def _generate_meeting_notes_jinja(title: str, context: dict) -> str:
        return _render_jinja_template(title, context, _JINJA_TEMPLATE_MEETING_NOTES)

    def _generate_project_proposal_jinja(title: str, context: dict) -> str:
        return _render_jinja_template(title, context, _JINJA_TEMPLATE_PROJECT_PROPOSAL)

    def _generate_code_review_jinja(title: str, context: dict) -> str:
        return _render_jinja_template(title, context, _JINJA_TEMPLATE_CODE_REVIEW)

    def render_docxtpl_with_template(title: str, context: dict, template_path: str) -> str:
        DocxTemplate = require_docxtpl()
        template_file = Path(template_path)
        if not template_file.exists():
            raise FileNotFoundError(f'Template file not found:\n{template_file}')
        doc = DocxTemplate(str(template_file))
        render_context = {'title': title, **context}
        doc.render(render_context)
        path = build_output_filename(title, 'docx')
        doc.save(str(path))
        return str(path)

    def _parse_experience(text: str) -> list[dict]:
        entries = []
        blocks = text.split('\n---\n')
        for block in blocks:
            lines = block.strip().split('\n')
            if not lines or not lines[0]: continue
            header_parts = [p.strip() for p in lines[0].split(';')]
            entries.append({
                'title': header_parts[0] if len(header_parts) > 0 else '',
                'company_info': ' | '.join([p for p in header_parts[1:] if p]),
                'description': '\n'.join(lines[1:]),
            })
        return entries

    def _parse_education(text: str) -> list[dict]:
        entries = []
        blocks = text.split('\n---\n')
        for block in blocks:
            parts = [p.strip() for p in block.strip().split(';')]
            entries.append({
                'degree': parts[0] if len(parts) > 0 else '',
                'school_info': ' | '.join([p for p in parts[1:] if p]),
            })
        return entries

    def _parse_budget(text: str) -> tuple[list[dict], float]:
        entries = []
        total_cost = 0.0
        for line in text.split('\n'):
            parts = [p.strip() for p in line.split(',', 2)]
            if len(parts) == 3:
                try:
                    cost = float(parts[2])
                    total_cost += cost
                    entries.append({'name': parts[0], 'desc': parts[1], 'cost': cost})
                except (ValueError, IndexError):
                    continue
        return entries, total_cost

    def _process_resume_context(context: dict):
        context['experience_entries'] = _parse_experience(context.get('experience_body', ''))
        context['education_entries'] = _parse_education(context.get('education_body', ''))

    def _process_project_proposal_context(context: dict):
        context['budget_entries'], context['budget_total'] = _parse_budget(context.get('budget_items', ''))

    _CONTEXT_PROCESSORS = {
        SUBJECT_RESUME: _process_resume_context,
        SUBJECT_PROJECT_PROPOSAL: _process_project_proposal_context,
    }

    _GENERATORS = {
        (FORMAT_PDF, SUBJECT_RESUME): _generate_resume_pdf,
        (FORMAT_PDF, SUBJECT_LETTER): _generate_letter_pdf,
        (FORMAT_PDF, SUBJECT_COVER_LETTER): _generate_cover_letter_pdf,
        (FORMAT_PDF, SUBJECT_MEETING_NOTES): _generate_meeting_notes_pdf,
        (FORMAT_PDF, SUBJECT_PROJECT_PROPOSAL): _generate_project_proposal_pdf,
        (FORMAT_PDF, SUBJECT_CODE_REVIEW): _generate_code_review_pdf,
        (FORMAT_PDF, 'generic'): _generate_generic_pdf,

        (FORMAT_WORD, SUBJECT_RESUME): _generate_resume_word,
        (FORMAT_WORD, SUBJECT_LETTER): _generate_letter_word,
        (FORMAT_WORD, SUBJECT_COVER_LETTER): _generate_cover_letter_word,
        (FORMAT_WORD, SUBJECT_MEETING_NOTES): _generate_meeting_notes_word,
        (FORMAT_WORD, SUBJECT_PROJECT_PROPOSAL): _generate_project_proposal_word,
        (FORMAT_WORD, SUBJECT_CODE_REVIEW): _generate_code_review_word,
        (FORMAT_WORD, 'generic'): _generate_generic_word,

        (FORMAT_JINJA, SUBJECT_RESUME): _generate_resume_jinja,
        (FORMAT_JINJA, SUBJECT_LETTER): _generate_letter_jinja,
        (FORMAT_JINJA, SUBJECT_COVER_LETTER): _generate_cover_letter_jinja,
        (FORMAT_JINJA, SUBJECT_MEETING_NOTES): _generate_meeting_notes_jinja,
        (FORMAT_JINJA, SUBJECT_PROJECT_PROPOSAL): _generate_project_proposal_jinja,
        (FORMAT_JINJA, SUBJECT_CODE_REVIEW): _generate_code_review_jinja,
        (FORMAT_JINJA, 'generic'): _generate_generic_jinja,
    }

    def generate_document(
        output_format: str,
        subject_key: str,
        use_docxtpl: bool,
        template_path: str,
        context_override: dict | None = None,
    ) -> str:
        if output_format not in set(OUTPUT_FORMATS):
            raise ValueError(f'Output format must be one of {sorted(OUTPUT_FORMATS)}')
        if subject_key not in set(SUBJECTS):
            raise ValueError(f'Subject must be one of {sorted(SUBJECTS)}')

        title, context = get_default_context_for_subject(subject_key)
        if context_override:
            context.update(context_override)

        # Pre-process context for rich structures if a processor is defined
        if processor := _CONTEXT_PROCESSORS.get(subject_key):
            processor(context)

        if use_docxtpl:
            if not template_path:
                raise ValueError("Please select a .docx template file or turn off 'Use docxtpl'.")
            return render_docxtpl_with_template(title, context, template_path)

        generator = _GENERATORS.get((output_format, subject_key)) or _GENERATORS.get((output_format, 'generic'))

        if not generator:
            raise ValueError(f"Unsupported output format or subject combination: {output_format}, {subject_key}")

        return generator(title, context)

    # ---------------- UI ----------------
    root = app.make_pop_ups_window(function=app.file_template_generator, title='Document Template Generator')

    # State variables
    output_format_var = tk.StringVar(value=FORMAT_WORD)
    subject_choice_var = tk.StringVar(value=SUBJECT_RESUME)
    use_docxtpl_var = tk.BooleanVar(value=False)
    template_path_var = tk.StringVar(value='')
    generated_path_var = tk.StringVar(value='')
    status_text_var = tk.StringVar(value='Select options and click Generate')

    # Choices
    output_format_options = OUTPUT_FORMATS
    subject_options = sorted(list(SUBJECTS))

    # Dynamic context fields state
    context_field_widgets: dict[str, tk.StringVar | tk.Text] = {}

    # Field names containing these substrings will use a multi-line Text widget
    MULTILINE_FIELD_SUBSTRINGS = ('_body', '_items', '_address', '_action', 'summary', 'description')

    # --- UI Helper Functions ---
    def _reset_generation_state(*_):
        if generated_path_var.get():
            generated_path_var.set('')
            status_text_var.set('Options changed. Click Generate to create a new file.')
            open_folder_button.config(state='disabled')
            open_file_button.config(state='disabled')
            copy_path_button.config(state='disabled')

    # --- Layout ---
    root.grid_columnconfigure(0, weight=1)
    root.grid_rowconfigure(0, weight=1)  # Make notebook expandable

    notebook = ttk.Notebook(root)
    notebook.grid(row=0, column=0, sticky='nsew', padx=8, pady=4)

    tab_settings = ttk.Frame(notebook)
    tab_content = ttk.Frame(notebook)
    notebook.add(tab_settings, text='Settings')
    notebook.add(tab_content, text='Content')

    # Make content tab expandable
    tab_content.grid_columnconfigure(0, weight=1)
    tab_content.grid_rowconfigure(0, weight=1)
    tab_settings.grid_columnconfigure(0, weight=1)

    # --- Frames ---
    main_options_lf = ttk.LabelFrame(tab_settings, text='Output Settings')
    main_options_lf.grid(row=0, column=0, sticky='ew', padx=8, pady=4)
    main_options_lf.grid_columnconfigure(3, weight=1)

    docxtpl_lf = ttk.LabelFrame(tab_settings, text='Docx Template')
    # This frame is hidden/shown by a checkbox

    # This frame will hold the canvas and scrollbar for the context fields
    context_container_lf = ttk.LabelFrame(tab_content, text='Placeholders (editable)')
    context_container_lf.grid(row=0, column=0, sticky='nsew', padx=8, pady=4)
    context_container_lf.grid_columnconfigure(0, weight=1)
    context_container_lf.grid_rowconfigure(0, weight=1)

    # Add a canvas and scrollbar for the context fields
    canvas = tk.Canvas(context_container_lf, borderwidth=0)
    scrollbar = ttk.Scrollbar(context_container_lf, orient="vertical", command=canvas.yview)
    context_lf = ttk.Frame(canvas)  # This frame holds the actual widgets

    canvas.configure(yscrollcommand=scrollbar.set)
    canvas.grid(row=0, column=0, sticky='nsew')
    scrollbar.grid(row=0, column=1, sticky='ns')
    canvas_frame = canvas.create_window((0, 0), window=context_lf, anchor="nw")

    def on_frame_configure(event):
        canvas.configure(scrollregion=canvas.bbox("all"))

    def on_canvas_configure(event):
        canvas.itemconfig(canvas_frame, width=event.width)

    context_lf.bind("<Configure>", on_frame_configure)
    canvas.bind('<Configure>', on_canvas_configure)

    actions_lf = ttk.LabelFrame(root, text='Actions')
    actions_lf.grid(row=1, column=0, sticky='ew', padx=8, pady=4)
    actions_lf.grid_columnconfigure(2, weight=1)

    status_label = ttk.Label(root, textvariable=status_text_var, anchor='w')
    status_label.grid(row=2, column=0, sticky='ew', padx=8, pady=(0, 8))

    # --- Widgets ---
    # Main Options
    ttk.Label(main_options_lf, text='Output format:').grid(row=0, column=0, sticky='w', padx=(4, 6))
    format_combo = ttk.Combobox(main_options_lf, textvariable=output_format_var, values=output_format_options,
                                state='readonly', width=10)
    format_combo.grid(row=0, column=1, sticky='w', padx=(0, 12))

    ttk.Label(main_options_lf, text='Subject:').grid(row=0, column=2, sticky='w', padx=(0, 6))
    subject_combo = ttk.Combobox(main_options_lf, textvariable=subject_choice_var, values=subject_options,
                                 state='readonly', width=18)
    subject_combo.grid(row=0, column=3, sticky='w')

    docxtpl_check = ttk.Checkbutton(main_options_lf, text='Use Custom .docx Template', variable=use_docxtpl_var)
    docxtpl_check.grid(row=1, column=0, columnspan=4, sticky='w', padx=4, pady=(6, 0))

    # Docx Template Frame
    ttk.Label(docxtpl_lf, text='Path:').grid(row=0, column=0, sticky='e', padx=(12, 6), pady=(2, 0))
    template_entry = ttk.Entry(docxtpl_lf, textvariable=template_path_var, width=36)
    template_entry.grid(row=0, column=1, sticky='we', pady=(2, 0))
    browse_template_btn = ttk.Button(docxtpl_lf, text='Browse...')
    browse_template_btn.grid(row=0, column=2, sticky='w', padx=(6, 4), pady=(2, 0))
    docxtpl_lf.grid_columnconfigure(1, weight=1)

    class CollapsiblePane(ttk.Frame):
        def __init__(self, parent, text="", expanded=True):
            super().__init__(parent)
            self.columnconfigure(0, weight=1)
            self._expanded = tk.BooleanVar(value=expanded)

            style = ttk.Style(self)
            style.configure('Collapsible.TButton', padding=0)
            self._button = ttk.Checkbutton(self, text=text, variable=self._expanded,
                                           command=self._toggle, style='Toolbutton')
            self._button.grid(row=0, column=0, sticky='w')

            self.container = ttk.Frame(self, padding=5)
            self.container.grid(row=1, column=0, sticky='nsew')
            self.container.columnconfigure(1, weight=1)

            if not expanded:
                self.container.grid_remove()

        def _toggle(self):
            if self._expanded.get():
                self.container.grid(row=1, column=0, sticky='nsew')
            else:
                self.container.grid_remove()

    def rebuild_context_fields(*_):
        _reset_generation_state()
        for w in context_lf.winfo_children():
            w.destroy()
        context_field_widgets.clear()

        _title, default_ctx = get_default_context_for_subject(subject_choice_var.get())

        # Group fields by headings
        groups = OrderedDict()
        current_group_key = 'general'
        groups[current_group_key] = {'heading': None, 'fields': OrderedDict()}

        for key, value in default_ctx.items():
            if key.endswith('_heading'):
                current_group_key = key
                groups[current_group_key] = {'heading': value, 'fields': OrderedDict()}
            else:
                groups[current_group_key]['fields'][key] = value

        # Build UI from groups
        pane_row = 0
        for group_key, group_data in groups.items():
            heading_text = group_data['heading']
            fields = group_data['fields']

            if not fields:
                continue

            # Use a default heading for the general section
            pane_title = heading_text if heading_text else "General Information"

            pane = CollapsiblePane(context_lf, text=pane_title, expanded=True)
            pane.grid(row=pane_row, column=0, sticky='ew', pady=(10, 0), padx=5)
            pane.columnconfigure(0, weight=1)
            container = pane.container
            pane_row += 1

            field_row = 0
            for key, default_value in fields.items():
                label = ttk.Label(container, text=key.replace('_', ' ').title() + ':')

                is_multiline = any(s in key for s in MULTILINE_FIELD_SUBSTRINGS)
                if is_multiline:
                    label.grid(row=field_row, column=0, sticky='ne', padx=4, pady=4)
                    widget = tk.Text(container, height=4, width=40, wrap=tk.WORD)
                    widget.insert('1.0', default_value)
                    widget.grid(row=field_row, column=1, sticky='we', padx=4, pady=2)

                    def placeholder_select_text(event, w=widget):
                        content = w.get('1.0', 'end-1c').strip()
                        if content.startswith('<') and content.endswith('>'):
                            w.after_idle(lambda: w.tag_add('sel', '1.0', 'end-1c'))
                            w.after_idle(lambda: w.mark_set('insert', 'end-1c'))

                    widget.bind("<FocusIn>", placeholder_select_text)
                    context_field_widgets[key] = widget
                else:
                    label.grid(row=field_row, column=0, sticky='e', padx=4, pady=2)
                    sv = tk.StringVar(value=default_value)
                    entry = ttk.Entry(container, textvariable=sv, width=40)
                    entry.grid(row=field_row, column=1, sticky='we', padx=4, pady=2)

                    def placeholder_select_entry(event, var=sv):
                        widget = event.widget
                        if var.get().startswith('<') and var.get().endswith('>'):
                            widget.after_idle(widget.select_range, 0, tk.END)
                            widget.after_idle(widget.icursor, tk.END)

                    entry.bind("<FocusIn>", placeholder_select_entry)
                    context_field_widgets[key] = sv
                field_row += 1
        context_lf.columnconfigure(0, weight=1)

    generate_button = ttk.Button(actions_lf, text='Generate')
    generate_button.grid(row=0, column=0, sticky='w', padx=4, pady=4)

    ttk.Label(actions_lf, text='Result path:').grid(row=0, column=1, sticky='e', padx=(12, 6))
    result_entry = ttk.Entry(actions_lf, textvariable=generated_path_var, state='readonly')
    result_entry.grid(row=0, column=2, sticky='ew')
    copy_path_button = ttk.Button(actions_lf, text='Copy Path', state='disabled')
    copy_path_button.grid(row=0, column=3, sticky='w', padx=(6, 0))

    open_folder_button = ttk.Button(actions_lf, text='Open Folder', state='disabled')
    open_folder_button.grid(row=1, column=0, sticky='w', padx=4, pady=4)
    open_file_button = ttk.Button(actions_lf, text='Open File', state='disabled')
    open_file_button.grid(row=1, column=1, columnspan=2, sticky='w')
    close_button = ttk.Button(actions_lf, text='Close')
    close_button.grid(row=1, column=3, sticky='e', padx=4, pady=4)

    actions_lf.grid_columnconfigure(2, weight=1)

    # ---------------- Callbacks ----------------
    def _open_path_in_explorer(path_to_open: str):
        try:
            if os.name == 'nt':
                os.startfile(path_to_open)
            elif sys.platform == 'darwin':
                subprocess.Popen(['open', path_to_open])
            else:
                subprocess.Popen(['xdg-open', path_to_open])
        except OSError as e:
            messagebox.showwarning('Error Opening Path', f'Could not open path:\n{path_to_open}\n\nError: {e}')
        except Exception as e:
            messagebox.showwarning('Error Opening Path', f'An unexpected error occurred:\n{e}')

    def _update_docxtpl_state(*_):
        if use_docxtpl_var.get():
            docxtpl_lf.grid(row=1, column=0, sticky='ew', padx=8, pady=4)
        else:
            docxtpl_lf.grid_forget()
        _reset_generation_state()

    def on_browse_template():
        path = filedialog.askopenfilename(
            title='Select .docx template',
            filetypes=[('Word Template', '*.docx'), ('All Files', '*.*')],
        )
        if path:
            template_path_var.set(path)
            use_docxtpl_var.set(True)

    def on_generate(event=None):
        fmt = output_format_var.get()
        subj = subject_choice_var.get()
        use_tpl = bool(use_docxtpl_var.get())
        tpl_path = template_path_var.get().strip()
        try:
            overrides = {}
            for key, var_or_widget in context_field_widgets.items():
                if isinstance(var_or_widget, tk.StringVar):
                    overrides[key] = var_or_widget.get()
                elif isinstance(var_or_widget, tk.Text):
                    overrides[key] = var_or_widget.get('1.0', 'end-1c')

            path = generate_document(fmt, subj, use_tpl, tpl_path, context_override=overrides)
            generated_path_var.set(str(path))
            status_text_var.set('Generated successfully.')
            open_folder_button.config(state=tk.NORMAL)
            open_file_button.config(state=tk.NORMAL)
            copy_path_button.config(state=tk.NORMAL)
            result_entry.configure(state='normal')
            result_entry.select_range(0, tk.END)
            result_entry.configure(state='readonly')
        except Exception as e:
            status_text_var.set('Generation failed.')
            messagebox.showerror('Generation error', str(e))

    def on_open_folder():
        p = generated_path_var.get()
        if p:
            folder = str(Path(p).resolve().parent)
            _open_path_in_explorer(folder)

    def on_open_file():
        p = generated_path_var.get()
        if p:
            file_path = str(Path(p).resolve())
            _open_path_in_explorer(file_path)

    def on_copy_path():
        p = generated_path_var.get()
        if not p:
            return
        try:
            root.clipboard_clear()
            root.clipboard_append(p)
            status_text_var.set('Path copied to clipboard.')
        except tk.TclError:
            status_text_var.set('Could not copy path to clipboard.')

    def on_close():
        try:
            root.destroy()
        except tk.TclError:
            pass

    # Wire callbacks
    browse_template_btn.configure(command=on_browse_template)
    generate_button.configure(command=on_generate)
    open_folder_button.configure(command=on_open_folder)
    open_file_button.configure(command=on_open_file)
    copy_path_button.configure(command=on_copy_path)
    close_button.configure(command=on_close)

    output_format_var.trace_add('write', _reset_generation_state)
    subject_choice_var.trace_add('write', rebuild_context_fields)
    use_docxtpl_var.trace_add('write', _update_docxtpl_state)

    root.bind('<Return>', on_generate)
    root.bind('<Escape>', lambda *_: on_close())

    try:
        if text_widget.tag_ranges('sel'):
            selection = text_widget.get('sel.first', 'sel.last')
            if selection and selection.strip().lower().endswith('.docx'):
                template_path_var.set(selection.strip())
                use_docxtpl_var.set(True)
    except Exception:
        pass

    # Set initial state
    rebuild_context_fields()
    _update_docxtpl_state()

    format_combo.focus_set()
